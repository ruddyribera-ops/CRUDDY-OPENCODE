# -*- coding: utf-8 -*-
"""Create GUIA-PARA-PROFESORES.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\Windows\.config\opencode\skills\education\weekly-planner-generator\GUIA-PARA-PROFESORES.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=RGBColor(0, 51, 102)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def para(doc, text, bold=False, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_checklist(doc):
    items = [
        "Carpeta MIS_PDC creada",
        "Archivos PDC organizados por materia",
        "Archivos nombrados correctamente (PDC_Materia_Grado_Nivel.docx)",
        "Carpeta de salida identificada",
        "Prompt copiado y rutas reemplazadas",
        "AI ejecutando la generacion",
        "Verificacion completada",
        "Archivos revisados",
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run("☐ " + item)
        run.font.size = Pt(10)

# ─── Build document ───
doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("GUIA PARA PROFESORES")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
run = subtitle.add_run("De PDC a Planificaciones Semanales")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

# ─── Section 1: Overview ───
heading(doc, "Overview", 1)
para(doc, "Este proceso convierte tus PDC (Planes de Desarrollo Curricular) en planificaciones semanales automaticamente. Tu AI generara 16-17 documentos por cada materia/grado.", size=11)

# What you need
heading(doc, "Lo que necesitas:", 2, color=RGBColor(0, 80, 0))
bullet(doc, "Archivos PDC organizados en carpetas")
bullet(doc, "Una carpeta de destino para las planificaciones")
bullet(doc, "Copiar y pegar un prompt en tu AI favorita")

# What AI does
heading(doc, "Lo que la AI hace sola:", 2, color=RGBColor(0, 80, 0))
bullet(doc, "Detecta automaticamente tus materias y grados")
bullet(doc, "Genera todos los documentos")
bullet(doc, "Verifica que todo coincida")

doc.add_paragraph()

# ─── Section 2: Paso 1 ───
heading(doc, "PASO 1: Organiza tus PDC", 1)

para(doc, "Estructura de carpetas requerida:", bold=True, size=11)

# Code block style
code_para = doc.add_paragraph()
code_para.paragraph_format.left_indent = Cm(0.5)
code_para.paragraph_format.space_after = Pt(4)
run = code_para.add_run("MIS_PDC/")
run.font.size = Pt(9)
run.font.name = 'Courier New'

for line in [
    "├── Tecnologia/",
    "│   ├── PRIMARIA/",
    "│   │   └── PDC_Tecnologia_1_Primaria.docx",
    "│   │   └── PDC_Tecnologia_2_Primaria.docx",
    "│   │   └── ... (un archivo por grado)",
    "│   └── SECUNDARIA/",
    "│       └── PDC_Tecnologia_1_Secundaria.docx",
    "├── Matematica/",
    "│   └── ...",
    "└── (todas las materias que ensenes)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.font.size = Pt(9)
    r.font.name = 'Courier New'

doc.add_paragraph()

# Naming rules
heading(doc, "Reglas para nombrar archivos PDC", 2)
para(doc, 'Formato: PDC_{Materia}_{Grado}_{Nivel}.docx', bold=True, size=11)

# Table: naming examples
tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Nivel", "Ejemplo"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

data = [
    ("Primaria", "PDC_Matematica_4_Primaria.docx"),
    ("Secundaria", "PDC_Historia_2_Secundaria.docx"),
]
for row_idx, (nivel, ejemplo) in enumerate(data):
    tbl.rows[row_idx + 1].cells[0].text = nivel
    tbl.rows[row_idx + 1].cells[1].text = ejemplo

# Set column widths
for row in tbl.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(9)

doc.add_paragraph()

# Grades table
heading(doc, "Grados y Niveles disponibles:", 2)
tbl2 = doc.add_table(rows=7, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["Grado", "Nivel", "Significado"]
for i, h in enumerate(headers2):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

grade_data = [
    ("1, 2, 3, 4, 5, 6", "Primaria", "1P, 2P, 3P, 4P, 5P, 6P"),
    ("1, 2, 3, 4, 5, 6", "Secundaria", "1S, 2S, 3S, 4S, 5S, 6S"),
]
for row_idx, (grado, nivel, sig) in enumerate(grade_data):
    tbl2.rows[row_idx + 1].cells[0].text = grado
    tbl2.rows[row_idx + 1].cells[1].text = nivel
    tbl2.rows[row_idx + 1].cells[2].text = sig

for row in tbl2.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(5)

doc.add_paragraph()

# ─── Section 3: Paso 2 ───
heading(doc, "PASO 2: Copia el Prompt para tu AI", 1)

para(doc, 'Abre tu AI favorita (ChatGPT, Claude, Gemini, etc.) y copia este prompt:', size=11)

# Prompt box
prompt_box = doc.add_paragraph()
prompt_box.paragraph_format.left_indent = Cm(0.5)
prompt_box.paragraph_format.space_after = Pt(4)
run = prompt_box.add_run("PROMPT (copia y pega en tu AI):")
run.font.bold = True
run.font.size = Pt(10)

# Prompt text in box
prompt_para = doc.add_paragraph()
prompt_para.paragraph_format.left_indent = Cm(0.5)
prompt_para.paragraph_format.space_after = Pt(4)
set_cell_bg_para = prompt_para
run = prompt_para.add_run(
    'Quiero generar planificaciones semanales automaticamente desde mis archivos PDC.\n\n'
    'Tengo mis archivos PDC organizados en esta carpeta:\n'
    '[RUTA COMPLETA A TU CARPETA MIS_PDC]\n\n'
    'Por ejemplo:\n'
    'C:/Users/MiNombre/Documents/MIS_PDC\n\n'
    'Genera planificaciones semanales para el Trimestre 2 (Semanas 15-30).\n\n'
    'Pasos que necesito:\n'
    '1. Lee todos los archivos PDC que estan en la carpeta\n'
    '2. Para cada materia y grado, genera 16 documentos de planificacion semanal (Semanas 15 a 30)\n'
    '3. Cada documento debe incluir:\n'
    '   - Numero de semana y fechas correctas\n'
    '   - Contenido de cada clase extraido del PDC\n'
    '   - Semana 23 y 24 marcadas como VACACIONES DE INVIERNO\n'
    '4. Guarda los archivos en esta carpeta de salida:\n'
    '[RUTA COMPLETA DONDE QUIERES LOS ARCHIVOS]\n'
    '5. Verifica que el contenido coincida con los PDC originales\n\n'
    'Confirma las rutas antes de empezar y dime cuantas materias y grados detectaste.'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph()

# How to get paths
heading(doc, "Como obtener las rutas:", 2)

path_items = [
    ('Windows:', 'Abre el explorador de archivos > Navega a tu carpeta MIS_PDC > Haz clic en la barra de direcciones > Copia todo (Ctrl+A, luego Ctrl+C)'),
    ('Mac:', 'Abre Finder > Navega a tu carpeta > Haz clic en el icono de carpeta en la barra superior > Arrastra al terminal para ver la ruta'),
]
for os_name, instructions in path_items:
    p = doc.add_paragraph()
    run = p.add_run(os_name + " ")
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(instructions)
    run2.font.size = Pt(10)

doc.add_paragraph()

# ─── Section 4: Paso 3 ───
heading(doc, "PASO 3: Ejecuta y Verifica", 1)

para(doc, "Lo que pasara automaticamente:", bold=True, size=11)

auto_items = [
    "La AI leera todos tus archivos PDC",
    "Detectara que materias y grados tienes",
    "Generara las planificaciones semana por semana",
    "Verificara que todo coincida",
]
for item in auto_items:
    bullet(doc, item)

doc.add_paragraph()

para(doc, "Confirmacion esperada:", bold=True, size=11)
confirm_box = doc.add_paragraph()
confirm_box.paragraph_format.left_indent = Cm(0.5)
run = confirm_box.add_run(
    "Detectado:\n"
    "- Tecnologia: 1P, 2P, 3P, 4P, 5P, 6P, 1S, 2S, 3S, 4S, 5S, 6S (12 grados)\n"
    "- Matematica: 1P, 2P, 3P, 4P, 5P, 6P (6 grados)\n"
    "- Historia: 1S, 2S, 3S, 4S (4 grados)\n\n"
    "Generando: 352 documentos...\n"
    "Verificando: todos los contenidos coinciden con PDC."
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph()

# ─── Section 5: Paso 4 ───
heading(doc, "PASO 4: Revisa los Archivos Generados", 1)

para(doc, "Ubicacion:", bold=True, size=11)
para(doc, "Tus planificaciones estaran en la carpeta de salida que indicaste.")

para(doc, "Cada archivo contiene:", bold=True, size=11)
bullet(doc, 'Semana y fechas (ej: "SEMANA 15 (11/05 - 15/05)")')
bullet(doc, "Unidad y contenido de cada clase")
bullet(doc, "Formato listo para imprimir")

doc.add_paragraph()

# ─── Section 6: Estructura PDC ───
heading(doc, "Estructura de un PDC", 1)

para(doc, "Para que la AI pueda leer tu PDC correctamente, debe tener esta estructura interna:", size=11)

pdc_items = [
    "TABLA 1 (Datos Referenciales): Distrito, Unidad Educativa, Nivel, etc.",
    "TABLA 2 (Desarrollo - 6 columnas): Objetivo, Contenidos, Momentos, Recursos, Periodos, Criterios",
    "La columna de Contenidos debe tener semanas en formato: 'Semana 15 (11/05 - 15/05) / Clase 1: Nombre de la clase'",
]
for item in pdc_items:
    bullet(doc, item)

doc.add_paragraph()

# ─── Section 7: FAQ ───
heading(doc, "Preguntas Frecuentes", 1)

faqs = [
    ("P: Que hago si la AI no detecta mis archivos?",
     "R: Verifica que la ruta sea correcta. Copia y pega la ruta exacta desde el explorador de archivos."),
    ("P: Mis horarios son diferentes a los estandares. Como los cambio?",
     'R: Inclucelos en el prompt personalizado. Ejemplo: "Mis horarios son: 1P: MIERCOLES 10:00-10:45, 2P: MARTES 14:00-14:45"'),
    ("P: Puedo generar solo para una materia?",
     'R: Si. En el prompt indica: "Solo genera para Tecnologia"'),
    ("P: Tengo mas de un trimestre. Como lo hago?",
     "R: Repite el proceso para cada carpeta de PDC (T1, T2, T3)."),
    ("P: Los archivos PDC necesitan tener semanas 23 y 24 como vacaciones?",
     "R: No. Si tu PDC no las menciona, la AI las agregara automaticamente."),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    run_q = p_q.add_run(q)
    run_q.font.bold = True
    run_q.font.size = Pt(10)

    p_a = doc.add_paragraph()
    run_a = p_a.add_run(a)
    run_a.font.size = Pt(10)
    p_a.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# ─── Section 8: Checklist ───
heading(doc, "Checklist de Preparacion", 1)
add_checklist(doc)

doc.add_paragraph()

# ─── Section 9: Tabla de semanas ───
heading(doc, "Datos Importantes - Trimestre 2", 1)

week_tbl = doc.add_table(rows=17, cols=2)
week_tbl.style = 'Table Grid'
week_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

week_headers = ["Semana", "Fechas"]
for i, h in enumerate(week_headers):
    cell = week_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

weeks = [
    (15, "11/05 - 15/05"), (16, "18/05 - 22/05"), (17, "25/05 - 29/05"),
    (18, "01/06 - 05/06"), (19, "08/06 - 12/06"), (20, "15/06 - 19/06"),
    (21, "22/06 - 26/06"), (22, "29/06 - 03/07"),
    (23, "VACACIONES 07/07 - 10/07"), (24, "VACACIONES 14/07 - 17/07"),
    (25, "20/07 - 24/07"), (26, "27/07 - 31/07"),
    (27, "04/08 - 08/08"), (28, "11/08 - 15/08"), (29, "18/08 - 22/08"), (30, "25/08 - 28/08"),
]

for row_idx, (week, dates) in enumerate(weeks):
    week_tbl.rows[row_idx + 1].cells[0].text = str(week)
    week_tbl.rows[row_idx + 1].cells[1].text = dates
    # Highlight vacaciones
    if week in [23, 24]:
        set_cell_bg(week_tbl.rows[row_idx + 1].cells[0], "FFD700")
        set_cell_bg(week_tbl.rows[row_idx + 1].cells[1], "FFD700")

for row in week_tbl.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(8)

doc.add_paragraph()

# ─── Footer ───
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(" weekly-planner-generator | Skill para Profesores | Bolivia MESCP ")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save
doc.save(OUT)
print(f"Guardado: {OUT}")