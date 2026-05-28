# -*- coding: utf-8 -*-
"""Create PDC guide DOCX - dead simple for teachers."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Temp\opencode\GUIA-CREAR-PDC.docx"

NAVY = RGBColor(0, 51, 102)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
ORANGE = RGBColor(200, 80, 0)
WHITE = RGBColor(255, 255, 255)
GREEN = RGBColor(0, 120, 0)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)


def shd_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t.upper())
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY


def body(t, sz=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(t)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.color.rgb = DARK
    return p


def step(num, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"PASO {num}: {title.upper()}")
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = ORANGE

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK


def copy_box(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    r = p.add_run(t)
    r.font.size = Pt(8.5)
    r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(40, 40, 40)


def tip(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TIP: " + t)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY


def ok_box(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F5E9')
    pPr.append(shd)
    r = p.add_run("  " + t)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = GREEN


def example_box(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8E1')
    pPr.append(shd)
    r = p.add_run(t)
    r.font.size = Pt(9)
    r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(80, 60, 0)


def div():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("─" * 75)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(200, 200, 200)


# ─── PORTADA ───
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(30)
r = t.add_run("GUIA RAPIDA")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = NAVY

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.paragraph_format.space_after = Pt(4)
r = s.add_run("Como crear un PDC")
r.font.size = Pt(15)
r.font.color.rgb = DARK

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.paragraph_format.space_after = Pt(30)
r = s2.add_run("Plan de Desarrollo Curricular - Sin conocimientos tecnicos")
r.font.size = Pt(10)
r.font.italic = True
r.font.color.rgb = GRAY

big = doc.add_paragraph()
big.alignment = WD_ALIGN_PARAGRAPH.CENTER
big.paragraph_format.space_after = Pt(6)
r = big.add_run("10 MINUTOS")
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = ORANGE

s3 = doc.add_paragraph()
s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
s3.paragraph_format.space_after = Pt(30)
r = s3.add_run("Tiempo para generar UN PDC (una materia, un grado)")
r.font.size = Pt(10)
r.font.color.rgb = GRAY

div()

# ─── QUE ES UN PDC ───
h1("Que es un PDC?")
body("PDC = Plan de Desarrollo Curricular.")
body("Es un documento obligatorio que todo profesor de Bolivia debe tener.")
body("Contiene:")
body("  - Las unidades que vas a enseñar este trimestre", sz=10)
body("  - Las clases de cada unidad", sz=10)
body("  - Como vas a evaluar", sz=10)
body("  - Que materiales vas a usar", sz=10)
doc.add_paragraph()
ok_box("Un PDC por materia y grado. Si enseñas Tecnologia en 6 grados, necesitas 6 PDC.")

h1("Los 4 archivos que necesitas")
body("Para que la AI cree un buen PDC, necesitas subirle 4 cosas:")
doc.add_paragraph()

tbl4 = doc.add_table(rows=5, cols=2)
tbl4.style = 'Table Grid'
files4 = [
    ("Archivo", "Que es"),
    ("1. PDC Trimestre 1", "Tus PDC del trim anterior. La AI los usa como referencia del contenido."),
    ("2. Plantilla del nuevo formato", "Un ejemplo de como debe quedar. Si no tienes, pidele a la AI que te muestre uno."),
    ("3. Contenido a enseņar", "Tus unidades y clases. Si no los tienes, la AI te los puede sugerir."),
    ("4. Calendario escolar", "Con las fechas de cada semana del Trimestre 2."),
]
for ri, (a, b) in enumerate(files4):
    tbl4.rows[ri].cells[0].text = a
    tbl4.rows[ri].cells[1].text = b
    if ri == 0:
        shd_cell(tbl4.rows[ri].cells[0], "003366")
        shd_cell(tbl4.rows[ri].cells[1], "003366")
        for c in tbl4.rows[ri].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = WHITE
    else:
        for c in tbl4.rows[ri].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    tbl4.rows[ri].cells[0].width = Cm(5)
    tbl4.rows[ri].cells[1].width = Cm(7)

tip("Si no tienes la plantilla, adjunta solo tus PDC del T1 y pidele a la AI: 'Generame un PDC siguiendo el formato MESCP de Bolivia'")
doc.add_paragraph()

# ─── FLUJO COMPLETO ───
h1("El flujo completo")
body("Mira esta guia primero para entender todo el proceso:", sz=10)
doc.add_paragraph()

# Table: flow
tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
flow = [
    ("PASO 1", "Crear PDC\n(con esta guia)"),
    ("PASO 2", "Crear Planificaciones Semanales\n(guia separada)"),
]
for ri, (s, d) in enumerate(flow):
    tbl.rows[ri].cells[0].text = s
    tbl.rows[ri].cells[1].text = d
    shd_cell(tbl.rows[ri].cells[0], "003366")
    for p in tbl.rows[ri].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = WHITE
    tbl.rows[ri].cells[0].width = Cm(3)
    tbl.rows[ri].cells[1].width = Cm(9)
    for p in tbl.rows[ri].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)

div()

# ─── 4 PASOS ───
h1("4 Pasos para crear tu PDC")

step(1, "Reune tus 4 archivos",
     "Necesitas tener estos archivos listos para subir al chat:\n\n"
     "  1. TUS PDC DEL TRIMESTRE 1  (como referencia del contenido)\n"
     "  2. UNA PLANTILLA DEL NUEVO FORMATO  (ejemplo de como debe quedar)\n"
     "  3. EL CONTENIDO A ENSEÑAR  (tus unidades y clases)\n"
     "  4. EL CALENDARIO ESCOLAR  (con las fechas del trimestre)\n\n"
     "Sube los 4 archivos al chat antes de enviar el prompt.")

step(2, "Copia y pega el texto de abajo",
     "Selecciona todo el texto del recuadro gris.\n"
     "Ctrl+A para seleccionar todo, Ctrl+C para copiar, Ctrl+V para pegar.\n\n"
     "IMPORTANTE: Antes de enviar, llena los datos de tu escuela.\n"
     "Busca [NOMBRE DE TU ESCUELA] y remplazalo con el nombre real.")

copy_box("""HOLA. NECESITO QUE ME CREES UN PDC (PLAN DE DESARROLLO CURRICULAR).

ARRIBA TE HE SUBIDO 4 ARCHIVOS:
1. MIS PDC DEL TRIMESTRE 1  (como referencia de contenido)
2. UNA PLANTILLA DEL NUEVO FORMATO  (como debe quedar el documento)
3. EL CONTENIDO A ENSEÑAR  (mis unidades y clases)
4. EL CALENDARIO ESCOLAR  (con las fechas del Trimestre 2)

MI ESCUELA:
- Nombre: [NOMBRE DE TU ESCUELA]
- nivel: [Primaria o Secundaria]
- Grado: [Numero de grado]
- Director/a: [NOMBRE DEL DIRECTOR]

MI MATERIA: [Nombre de la materia, ejemplo: Tecnologia]
YO SOY EL PROFESOR: [TU NOMBRE]

INSTRUCCIONES:
1. REVISA que hayas recibido los 4 archivos que adjunte
2. USA el contenido de mis PDC del Trimestre 1 como referencia
3. USA la plantilla que adjunte como el formato exacto a seguir
4. USA el calendario escolar para poner las fechas correctas
5. GENERA el PDC siguiendo el formato MESCP de Bolivia
6. Incluye las 6 columnas: Objetivo, Contenidos, Momentos, Recursos, Periodos, Criterios
7. Incluye: Adaptaciones Curriculares, Inteligencias Multiples, Productos
8. Usa las fechas del Trimestre 2 (Semana 15 a 30, incluyendo Vacaciones de Invierno)
9. Entregame el documento en formato DOCX (Word) para descargar

ANTES DE GENERAR, CONFIRMAME:
- Que recibiste los 4 archivos
- Que entendiste el formato de la plantilla
- Cuantos PDC vas a generar

GRACIAS.""")

tip("Si no sabes que unidades escribir, pidele a la AI que te sugiera basandose en el nombre de tu materia.")

step(3, "Llena tus datos antes de enviar",
     "Busca en el texto que pegaste:\n"
     "  [NOMBRE DE TU ESCUELA]  -  remplazalo con el nombre\n"
     "  [Primaria o Secundaria]  -  cual es tu nivel\n"
     "  [Numero de grado]  -  ejemplo: 4, 5, 2\n"
     "  [NOMBRE DEL DIRECTOR]  -  nombre del director\n"
     "  [Nombre de la materia]  -  Tecnologia, Matematica, etc.\n"
     "  [TU NOMBRE]  -  tu nombre completo\n\n"
     "Tambien escribe las unidades que vas a enseņar.")

step(4, "Envia y descarga",
     "La AI te va a entregar el documento DOCX para descargar.\n"
     "Descarga el archivo Word.\n"
     "Revisa que este completo y descargalo.")

div()

# ─── EJEMPLO REAL ───
h1("Ejemplo de unidades (como escribirlas)")

body("Este es un ejemplo real de como escribir las unidades para Tecnologia 5to Primaria:")

example_box("""UNIDAD 6: PYTHON BASICO EN ENTORNOS VISUALES
CLASE 1: Repaso de conceptos basicos
CLASE 2: Variables y tipos de datos
CLASE 3: Entrada y salida de datos
CLASE 4: Condicionales if-else
CLASE 5: Proyecto con condicionales

UNIDAD 7: FUNCIONES Y ARCHIVOS
CLASE 1: Creando funciones
CLASE 2: Parametros y valores de retorno
CLASE 3: Lectura de archivos
CLASE 4: Escritura de archivos
CLASE 5: Proyecto integrador

UNIDAD 8: PROYECTO FINAL
CLASE 1: Planificacion del proyecto
CLASE 2: Desarrollo del proyecto
CLASE 3: Presentacion y evaluacion""")

tip("Puedes escribir mas o menos clases por unidad. Solo sigue el formato: CLASE numero: nombre")

div()

# ─── COMBINAR GUIAS ───
h1("Ya tienes PDC? Haz las planificaciones semanales")

body("Si ya tienes tus PDC creados, la otra guia te ayuda a generar")
body("las planificaciones semanales automaticamente.")

ok_box("PDC + esta guia  =  Planificaciones Semanales")
body("Descarga la guia: 'GUIA-PARA-PROFESORES.docx'")

div()

# ─── FAQ ───
h1("Preguntas Frecuentes")

faqs = [
    ("No tengo la plantilla del nuevo formato. Que hago?",
     "Sube solo tus PDC del Trimestre 1 y pidele a la AI: 'Generame un PDC en formato MESCP de Bolivia basandote en mi contenido del T1'."),
    ("No tengo PDC del Trimestre 1. Que hago?",
     "Sube el contenido que vas a enseņar y el calendario. La AI puede crear el PDC desde cero."),
    ("Cuantas unidades necesito?",
     "Normalmente 4 a 8 unidades por trimestre. Una unidad por cada tema principal que enseñas."),
    ("Cuantas clases por unidad?",
     "Normalmente 3 a 6 clases. Una clase puede ser una o dos sesiones de 45 minutos."),
    ("Puedo pedirle a la AI que me sugiera las unidades?",
     "Si! Escribe: 'Sugiereme unidades para Tecnologia de 5to Primaria, Trimestre 2' y la AI te da ideas."),
    ("Puedo crear PDC para varias materias a la vez?",
     "Es mejor hacer uno a la vez. Crea un chat nuevo por cada materia/grado."),
    ("EI documento que me da la AI es editable?",
     "Si. Es un archivo Word (.docx). Lo puedes abrir, modificar e imprimir."),
]

for q, a in faqs:
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(6)
    pq.paragraph_format.space_after = Pt(1)
    rq = pq.add_run("P: " + q)
    rq.font.size = Pt(10)
    rq.font.bold = True
    rq.font.color.rgb = DARK

    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_after = Pt(6)
    ra = pa.add_run("R: " + a)
    ra.font.size = Pt(10)
    ra.font.color.rgb = GRAY

div()

# ─── FOOTER ───
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.paragraph_format.space_before = Pt(16)
r = f.add_run("Esta guia funciona con cualquier AI: ChatGPT, Gemini, Claude, Grok...")
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = GRAY

f2 = doc.add_paragraph()
f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = f2.add_run("Solo necesitas: nombre de tu materia + unidades + 10 minutos")
r2.font.size = Pt(9)
r2.font.bold = True
r2.font.color.rgb = NAVY

doc.save(OUT)
print(f"Guardado: {OUT}")