# -*- coding: utf-8 -*-
"""Create simple, beautiful guide DOCX for teachers."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Windows\Desktop\GUIA-PARA-PROFESORES.docx"

# Colors
NAVY = RGBColor(0, 51, 102)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
ORANGE = RGBColor(200, 100, 0)
WHITE = RGBColor(255, 255, 255)

doc = Document()

# Page
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def body(text, size=11, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = DARK
    return p

def step(num, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"PASO {num}: ")
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = ORANGE
    r2 = p.add_run(title.upper())
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(8)
    r3 = p2.add_run(desc)
    r3.font.size = Pt(10)
    r3.font.color.rgb = DARK

def important_box(text):
    """Shaded box for important text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = ORANGE
    return p

def copy_block(text):
    """Code-style block for copy-paste."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("TIP: " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(200, 200, 200)

# ─── TITLE PAGE ───
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(40)
title.paragraph_format.space_after = Pt(10)
run = title.add_run("GUIA RAPIDA")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
run = sub.add_run("De PDC a Planificaciones Semanales")
run.font.size = Pt(16)
run.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(40)
run = sub2.add_run("Para profesores - Sin conocimientos tecnicos")
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.italic = True

# Big number
big = doc.add_paragraph()
big.alignment = WD_ALIGN_PARAGRAPH.CENTER
big.paragraph_format.space_after = Pt(40)
run = big.add_run("5 MINUTOS")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = ORANGE

body("Tiempo estimado para generar todas tus planificaciones", 10, indent=0, bold=False)
body("del Trimestre 2.", 10, bold=False)

doc.add_paragraph()
divider()

# ─── OVERVIEW ───
h1("Lo que vamos a hacer")
body("Tus archivos PDC se van a convertir en planificaciones semanales automaticamente.")
body("Tu solo pegas un texto en ChatGPT o Gemini, subes tus archivos, y listo.")
body("No necesitas saber programacion ni nada tecnico.", 11, bold=True)

doc.add_paragraph()

# ─── WHAT YOU NEED ───
h1("Lo que necesitas")
body("Solo necesitas UNA COSA: tus archivos PDC.")
body("El AI va a generar las planificaciones usando el contenido de tus PDC.", bold=True)
doc.add_paragraph()

# What to upload
h2("Archivos que debes subir al AI:")

upload_items = [
    ("Tus archivos PDC",
     "Un archivo por materia y grado.\n"
     "Ejemplo: si ensenas Tecnologia en 6 grados, tendras 6 archivos PDC.\n"
     "Si ensenas Tecnologia y Matematicas en 4 grados, tendras 8 archivos."),
    ("El texto del recuadro negro",
     "Lo copias y pegas directamente en el chat. No es un archivo,\n"
     "es solo texto que le das al AI como instruccion."),
    ("Las rutas de tus carpetas",
     "Le dices al AI donde estan tus archivos y donde guardar el resultado.\n"
     "Ejemplo: C:/Users/Maria/Documents/MIS_PDC"),
]

for title, desc in upload_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("  " + title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK

body("No necesitas otros documentos. El AI lee tus PDC y genera todo.", bold=True)
body("Las fechas del calendario ya estan incluidas en el texto que vas a copiar.")
doc.add_paragraph()

divider()

# ─── STEPS ───
h1("4 Pasos Simples")

step(1, "Sube tus archivos PDC al chat",
     'Abre ChatGPT o Gemini.\n'
     'Busca el icono de adjuntar/clip en el chat y sube tus archivos PDC.\n\n'
     'Como nombrarlos (muy importante):\n'
     '  Tecnologia_1_Primaria.docx\n'
     '  Matematica_4_Secundaria.docx\n'
     '  Historia_2_Secundaria.docx\n\n'
     'El formato es:  Materia_Grado_Nivel.docx')

step(2, "Pega el texto de abajo en el chat",
     'Selecciona todo el texto del recuadro negro.\n'
     'Haz Ctrl+A para seleccionar todo, luego Ctrl+C para copiar.\n'
     'Pegalo en el chat con Ctrl+V.\n\n'
     'ANTES DE ENVIAR: asegurate de que tus archivos PDC esten subidos\n'
     '(Paso 1). El texto dice "archivos adjuntos" - subelos primero.')

copy_block("""
HOLA. NECESITO QUE GENERES MIS PLANIFICACIONES SEMANALES.

TENGO ARCHIVOS PDC (Planes de Desarrollo Curricular) Y QUIERO QUE LOS
CONVIERTAS EN PLANIFICACIONES SEMANALES PARA TODO EL TRIMESTRE.

YA TE HE ADJUNTADO MIS ARCHIVOS PDC. REVISA QUE LOS HAYAS RECIBIDO.

MI CARPETA DE SALIDA DONDE GUARDAR:
[RUTA DONDE QUIERES LOS ARCHIVOS - ejemplo: C:/Users/Maria/Documents/PLANIFICACIONES]

MI ESCUELA:
- Nombre: [NOMBRE DE TU ESCUELA]
- Trimestre: 2 (11 de mayo al 28 de agosto de 2026)
- Director/a: [NOMBRE DEL DIRECTOR]

POR FAVOR:

1. REVISA que hayas recibido mis archivos PDC adjuntos
2. CONFIRMAME que materias y grados detectas en ellos
3. GENERA 16 documentos de planificacion semanal para cada materia/grado
   (Semanas 15 a 30)
4. CADA documento debe tener:
   - Numero de semana y fechas correctas (usa el calendario del T2)
   - Contenido de cada clase (extraido del PDC)
   - Semanas 23 y 24 marcadas como VACACIONES DE INVIERNO
5. GUARDA los archivos en la carpeta de salida que indique
6. VERIFICA que el contenido coincida con mis PDC originales

CONFIRMAME antes de empezar:
- Cuantos archivos PDC recibiste
- Cuantos documentos vas a generar en total
- En que carpeta los vas a guardar

GRACIAS.
""")

tip("Para saber la ruta de salida: abre la carpeta donde quieres los archivos, haz clic en la barra de arriba, copiala con Ctrl+A y Ctrl+C")

step(3, "Espera y confirma",
     'La AI va a revisar tus archivos y te va a decir:\n'
     '  - Cuantos archivos PDC recibio\n'
     '  - Cuantos documentos va a generar\n'
     '  - En que carpeta los va a guardar\n\n'
     'Confirma que esta todo bien y dime "SI, GENERALOS".')

step(4, "Descarga los archivos",
     'Cuando termine de generar, te va a dar los archivos.\n'
     'Descarga el archivo ZIP que te den y descomprimelo.\n'
     'Ahi estan todas tus planificaciones listas para imprimir.')

divider()

# ─── WEEK TABLE ───
h1("Fechas del Trimestre 2")
body("Referencia rapida para que sepas que fechas corresponden a cada semana:")

table = doc.add_table(rows=17, cols=2)
table.style = 'Table Grid'

# Header
hdr = table.rows[0].cells
hdr[0].text = "Semana"
hdr[1].text = "Fechas"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = WHITE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '003366')
    tcPr.append(shd)

weeks_data = [
    ("15", "11/05 - 15/05"), ("16", "18/05 - 22/05"), ("17", "25/05 - 29/05"),
    ("18", "01/06 - 05/06"), ("19", "08/06 - 12/06"), ("20", "15/06 - 19/06"),
    ("21", "22/06 - 26/06"), ("22", "29/06 - 03/07"),
    ("23", "VACACIONES 07/07 - 10/07"), ("24", "VACACIONES 14/07 - 17/07"),
    ("25", "20/07 - 24/07"), ("26", "27/07 - 31/07"),
    ("27", "04/08 - 08/08"), ("28", "11/08 - 15/08"), ("29", "18/08 - 22/08"), ("30", "25/08 - 28/08"),
]

for row_idx, (week, dates) in enumerate(weeks_data):
    row = table.rows[row_idx + 1]
    row.cells[0].text = week
    row.cells[1].text = dates
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = DARK
        if week in ["23", "24"]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'FFD700')
            tcPr.append(shd)

# Widths
for row in table.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(10)

divider()

# ─── FAQ ───
h1("Preguntas Frecuentes")

faqs = [
    ("No tengo la barra de direcciones en Mac.",
     "En Finder, ve a la carpeta y presiona Comando + Opcion + C. Aparecera la ruta completa."),
    ("Mis archivos no se llaman como el ejemplo.",
     "Esta bien. Solo renombrarlos al formato: Materia_Grado_Nivel.docx"),
    ("Tengo solo una materia, funciona igual?",
     "Si. Funciona con cualquier cantidad de archivos."),
    ("Los archivos me los da en partes, no todos juntos?",
     "Normal. ChatGPT y Gemini tienen limite de descarga. Pide 'generame un ZIP' o 'descargame todo en un archivo'."),
    ("Las fechas de los documentos estan bien?",
     "Si la AI leyó bien tu PDC, si. Por eso la verificacion automatica es importante."),
]

for q, a in faqs:
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(8)
    pq.paragraph_format.space_after = Pt(2)
    rq = pq.add_run("P: " + q)
    rq.font.size = Pt(10)
    rq.font.bold = True
    rq.font.color.rgb = DARK

    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_after = Pt(6)
    ra = pa.add_run("R: " + a)
    ra.font.size = Pt(10)
    ra.font.color.rgb = GRAY

divider()

# ─── FOOTER ───
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_before = Pt(20)
run = footer.add_run("Esta guia funciona con cualquier AI: ChatGPT, Gemini, Claude, Grok, etc.")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.italic = True

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run("Solo necesitas tus archivos PDC organizados y 5 minutos.")
run2.font.size = Pt(9)
run2.font.bold = True
run2.font.color.rgb = NAVY

# Save
doc.save(OUT)
print(f"Guardado: {OUT}")