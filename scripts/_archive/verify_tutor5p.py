# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os, re

BASE_PDC = r'C:\Users\Windows\Desktop\PDC - T2 - 2026\Tutor 5p'
WEEKLY = r'C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\5TO PRIMARIA _ TUTOR\5P'

def parse_pdc_weeks(path):
    doc = Document(path)
    t = doc.tables[1]
    content = t.rows[1].cells[1].text

    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]

    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]

    def find_unit(wpos):
        for i in range(len(units)-1, -1, -1):
            if units[i][0] < wpos:
                return 'UNIDAD %s: %s' % (units[i][1], units[i][2])
        return ''

    result = {}
    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        next_pos = weeks_raw[idx+1][0] if idx+1 < len(weeks_raw) else len(content)
        raw = content[wpos:next_pos]
        lines = [l.strip() for l in raw.strip().split('\n')]
        content_lines = [l for l in lines[1:] if l and not re.match(r'^[\*]+', l) and not re.match(r'^UNIDAD', l)]
        result[wnum] = {
            'unit': find_unit(wpos),
            'content': '\n'.join(content_lines),
            'header': lines[0]
        }
    return result


# Slot -> (time_slot_idx, day, table_row, table_col)
# We need to know where each subject appears
# TIME_SLOTS from generator:
# 0: 07:55 LUNES=LENGUAJE, MARTES=MATEMATICAS, MIERCOLES=MATEMATICAS, JUEVES=MUSICA, VIERNES=MATEMATICAS
# 1: 08:40 LUNES=ED.FISICA, MARTES=MATEMATICAS, MIERCOLES=MATEMATICAS, JUEVES=SCIENCE, VIERNES=MATEMATICAS
# 2: 09:25 LUNES=INGLES ... all days=INGLES
# 3: 10:10 SNACK TIME
# 4: 10:30 LUNES=LENGUAJE, MARTES=ED.FISICA, MIERCOLES=SOCIALES, JUEVES=PORTUGUES, VIERNES=LENGUAJE
# 5: 11:15 LUNES=LENGUAJE, MARTES=LENGUAJE, MIERCOLES=SOCIALES, JUEVES=MATEMATICAS, VIERNES=LENGUAJE
# 6: 12:00 ACTIVE PAUSE
# 7: 12:15 LUNES=SCIENCE, MARTES=TECNOLOGIA, MIERCOLES=ARTES, JUEVES=ARTES, VIERNES=HABILIDADES

# Content rows: 2, 4, 6, 8, 10, 12, 14, 16 (time slot rows: 1, 3, 5, 7, 9, 11, 13, 15)
slot_map = {
    'LENGUAJE':     [(0, 'LUNES', 2, 1), (4, 'LUNES', 10, 1), (5, 'LUNES', 12, 1), (4, 'MARTES', 10, 2), (5, 'MARTES', 12, 2), (4, 'VIERNES', 10, 5), (5, 'VIERNES', 12, 5)],
    'MATEMATICAS':  [(0, 'MARTES', 2, 2), (0, 'VIERNES', 2, 5), (1, 'MARTES', 4, 2), (1, 'VIERNES', 4, 5), (5, 'JUEVES', 12, 4)],
    'SOCIALES':     [(4, 'MIERCOLES', 10, 3), (5, 'MIERCOLES', 12, 3)],
    'TECNOLOGIA':   [(7, 'MARTES', 16, 2)],
    'SCIENCE':      [(1, 'JUEVES', 4, 4), (7, 'LUNES', 16, 1)],
    'ED.FISICA':    [(0, 'LUNES', 2, 1), (1, 'LUNES', 4, 1), (4, 'MARTES', 10, 2)],
    'MUSICA':       [(0, 'JUEVES', 2, 4)],
    'PORTUGUES':    [(4, 'JUEVES', 10, 4)],
    'ARTES':        [(7, 'MIERCOLES', 16, 3), (7, 'JUEVES', 16, 4)],
    'HABILIDADES':  [(7, 'VIERNES', 16, 5)],
    'INGLES':       [(2, 'LUNES', 6, 1), (2, 'MARTES', 6, 2), (2, 'MIERCOLES', 6, 3), (2, 'JUEVES', 6, 4), (2, 'VIERNES', 6, 5)],
}

subjects = {
    'Lenguaje':    ('PDC_Lenguaje_5_Primaria.docx', 'LENGUAJE'),
    'Matematicas': ('PDC_Matematica_5_Primaria.docx', 'MATEMATICAS'),
    'Sociales':    ('PDC_Sociales_5_Primaria.docx', 'SOCIALES'),
    'Tecnologia':  ('PDC_Tecnologia_5_Primaria.docx', 'TECNOLOGIA'),
    'Valores':     ('PDC_Valores_5_Primaria.docx', None),  # Valores is not in the regular schedule
}

print('=' * 70)
print('TUTOR 5P VERIFICATION: PDC vs GENERATED WEEKLY')
print('=' * 70)

all_ok = True
errors = []

pdc_data = {}
for subj, (fname, _) in subjects.items():
    if fname:
        path = os.path.join(BASE_PDC, fname)
        if os.path.exists(path):
            pdc_data[subj] = parse_pdc_weeks(path)
            print('Loaded PDC: %s (%d weeks)' % (subj, len(pdc_data[subj])))
        else:
            print('MISSING PDC: %s' % path)

print()

for week in range(15, 31):
    weekly_path = os.path.join(WEEKLY, 'SEMANA %d - 5P.docx' % week)
    if not os.path.exists(weekly_path):
        errors.append('SEMANA %d: FILE MISSING' % week)
        continue

    doc = Document(weekly_path)
    t = doc.tables[0]

    for subj, (fname, subj_key) in subjects.items():
        if not fname or not subj_key:
            continue

        if subj_key not in slot_map:
            continue

        for slot_idx, day, row, col in slot_map[subj_key]:
            gen_text = t.rows[row].cells[col].text.strip() if row < len(t.rows) and col < len(t.rows[row].cells) else ''

            if week in [23, 24]:
                # Vacaciones - all content cells should say vacaciones
                if 'VACACIONES' not in gen_text and gen_text != '':
                    errors.append('SEMANA %d %s %s: Expected VACACIONES, got: %s' % (week, subj, day, gen_text[:50]))
                    all_ok = False
                continue

            if week not in pdc_data.get(subj, {}):
                # Not in PDC for this week
                continue

            pdc_week = pdc_data[subj][week]
            pdc_content = pdc_week['content']
            pdc_unit = pdc_week['unit']

            # Check content
            pdc_first = pdc_content.split('\n')[0][:30]

            if pdc_first and pdc_first not in gen_text:
                errors.append('SEMANA %d %s %s: CONTENT MISMATCH' % (week, subj, day))
                errors.append('  PDC: %s' % pdc_content[:80])
                errors.append('  GEN: %s' % gen_text[:80])
                all_ok = False

            # Check unit
            if pdc_unit and pdc_unit not in gen_text and pdc_unit.replace('UNIDAD ', '') not in gen_text:
                errors.append('SEMANA %d %s %s: UNIT MISSING' % (week, subj, day))
                errors.append('  Expected: %s' % pdc_unit)
                errors.append('  Got:      %s' % gen_text[:100])
                all_ok = False

if all_ok and not errors:
    print('ALL CHECKS PASSED')
else:
    print('\nERRORS (%d):' % len(errors))
    for e in errors[:30]:
        print('  ' + e)
    if len(errors) > 30:
        print('  ... and %d more' % (len(errors) - 30))

print('\nTotal errors:', len(errors))