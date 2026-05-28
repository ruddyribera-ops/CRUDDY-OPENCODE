# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os, re

BASE_PDC = r'C:\Users\Windows\Desktop\PDC - T2 - 2026\Tecnologia'
BASE_WEEKLY = r'C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\TECNOLOGIA'

# Grade info
GRADES = [
    ('1P', 'PRIMARIA', '1_Primaria'),
    ('2P', 'PRIMARIA', '2_Primaria'),
    ('3P', 'PRIMARIA', '3_Primaria'),
    ('4P', 'PRIMARIA', '4_Primaria'),
    ('6P', 'PRIMARIA', '6_Primaria'),
    ('1S', 'SECUNDARIA', '1_Secundaria'),
    ('2S', 'SECUNDARIA', '2_Secundaria'),
    ('3S', 'SECUNDARIA', '3_Secundaria'),
    ('4S', 'SECUNDARIA', '4_Secundaria'),
    ('5S', 'SECUNDARIA', '5_Secundaria'),
    ('6S', 'SECUNDARIA', '6_Secundaria'),
]

def parse_pdc_weeks(path):
    doc = Document(path)
    t = doc.tables[1]
    content = t.rows[1].cells[1].text

    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]

    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]

    def find_unit(wpos):
        for i in range(len(units)-1, -1, -1):
            if units[i][0] < wpos:
                return 'UNIDAD %s: %s' % (units[i][1], units[i][2])
        return ''

    result = {}
    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        next_pos = weeks_raw[idx+1][0] if idx+1 < len(weeks_raw) else len(content)
        raw = content[wpos:next_pos]
        lines = [l.strip() for l in raw.strip().split('\n')]
        content_lines = [l for l in lines[1:] if l and not re.match(r'^[\*]+', l) and not re.match(r'^UNIDAD', l)]
        result[wnum] = {
            'unit': find_unit(wpos),
            'content': '\n'.join(content_lines),
            'header': lines[0]
        }
    return result


print('=' * 70)
print('VERIFICATION: PDC vs GENERATED WEEKLY (Tecnologia standalone)')
print('=' * 70)

all_ok = True
errors = []

for grade, nivel, pdc_name in GRADES:
    pdc_path = os.path.join(BASE_PDC, nivel, 'PDC_Tecnologia_%s.docx' % pdc_name)
    weekly_dir = os.path.join(BASE_WEEKLY, nivel, grade)

    pdc_data = parse_pdc_weeks(pdc_path)

    grade_ok = True
    for week in range(15, 31):
        weekly_path = os.path.join(weekly_dir, 'SEMANA %d - %s.docx' % (week, grade))

        if not os.path.exists(weekly_path):
            errors.append('%s SEMANA %d: FILE MISSING' % (grade, week))
            grade_ok = False
            continue

        doc = Document(weekly_path)
        t = doc.tables[0]

        # Get content from generated doc (row 2, col 1)
        if week in [23, 24]:
            gen_content = t.rows[2].cells[0].text
        else:
            gen_content = t.rows[2].cells[1].text if t.rows[2].cells[1].text else t.rows[2].cells[0].text

        if week not in pdc_data:
            # Vacaciones or missing
            if week in [23, 24]:
                if 'VACACIONES' not in gen_content:
                    errors.append('%s SEMANA %d: Expected VACACIONES, got: %s' % (grade, week, gen_content[:50]))
                    grade_ok = False
            continue

        pdc_week = pdc_data[week]
        pdc_unit = pdc_week['unit']
        pdc_content = pdc_week['content']
        pdc_header = pdc_week['header']

        # Check: does generated content contain the PDC content?
        # The generated format is: "UNIDAD: X\nY" or just "Y"
        gen_clean = gen_content.replace('UNIDAD: ', '').replace('CONTENIDO: ', '').replace('\n', ' ')

        # PDC content might span multiple lines
        pdc_clean = pdc_content.strip()

        if pdc_clean not in gen_clean and pdc_clean.split('\n')[0] not in gen_clean:
            # Try more lenient check
            pdc_first_line = pdc_clean.split('\n')[0]
            if pdc_first_line[:30] not in gen_clean[:50]:
                errors.append('%s SEMANA %d: CONTENT MISMATCH' % (grade, week))
                errors.append('  PDC:     %s' % pdc_clean[:80])
                errors.append('  GEN:     %s' % gen_clean[:80])
                grade_ok = False

        # Check unit header
        if pdc_unit and pdc_unit not in gen_content:
            errors.append('%s SEMANA %d: UNIT MISSING' % (grade, week))
            errors.append('  Expected unit: %s' % pdc_unit)
            errors.append('  In content:    %s' % gen_content[:100])
            grade_ok = False

    status = 'OK' if grade_ok else 'ERRORS'
    print('  %s: %s' % (grade, status))

if errors:
    print('\nERRORS FOUND (%d):' % len(errors))
    for e in errors:
        print('  ' + e)
else:
    print('\nAll checks passed!')

print('\nTotal errors:', len(errors))