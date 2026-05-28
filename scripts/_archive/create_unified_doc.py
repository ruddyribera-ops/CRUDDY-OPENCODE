from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
sys.stdout.reconfigure(encoding='utf-8')

def set_heading_style(paragraph, level=1):
    """Apply heading style"""
    paragraph.style = f'Heading {level}'

def add_bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

# All content data
cursos = {
    "1RO DE PRIMARIA": [
        ("UNIDAD 1 – TECNOLOGÍA Y EVOLUCIÓN", [
            "CLASE 2.1 – TECNOLOGÍA",
            "CLASE 2.2 – EL AVANCE DE LA TECNOLOGÍA",
            "CLASE 2.3 – LAS TICS",
            "CLASE 2.4 – INFORMACIÓN PÚBLICA Y PRIVADA",
        ]),
        ("UNIDAD 2 – USO DEL TECLADO Y DEL MOUSE", [
            "CLASE 1.1 – USO DEL MOUSE",
            "CLASE 1.2 – USO DEL MOUSE SIN G COMPRIS",
            "CLASE 1.3 – USO DEL MOUSE PARTE 2",
            "CLASE 1.4 – USO DEL TECLADO",
            "CLASE 1.5 – USO DEL TECLADO SIN G COMPRIS",
            "CLASE 1.6 – USO DEL TECLADO PARTE 2",
        ]),
        ("UNIDAD 3 – ALGORITMOS", [
            "CLASE 3.1 – QUÉ ES UN ALGORITMO",
            "CLASE 3.2 – ALGORITMOS",
            "CLASE 3.3 – ALGORITMOS Y FLECHAS",
            "CLASE 3.4 – PRACTICAMOS ALGORITMOS",
        ]),
        ("UNIDAD 4 – BLOQUES DE DIRECCIÓN", [
            "CLASE 4.1 – QUÉ ES PILAS BLOQUES",
            "CLASE 4.2 – PROGRAMACIÓN CON PILAS BLOQUES",
            "CLASE 4.3 – LA MULITA Y SU ENSALADA",
            "CLASE 4.4 – LITA Y DUBA",
        ]),
        ("UNIDAD 5 – PATRONES", [
            "CLASE 5.1 – QUÉ SON LOS PATRONES",
            "CLASE 5.2 – CODIFICANDO",
            "CLASE 5.3 – CÓDIGO MURCIÉLAGO",
        ]),
        ("UNIDAD 6 – BUCLES", [
            "CLASE 6.1 – QUÉ ES UN BUCLE",
            "CLASE 6.2 – BUCLES EN PROGRAMACIÓN",
            "CLASE 6.3 – PRACTICAMOS BUCLES",
            "CLASE 6.4 – SEGUIMOS PRACTICANDO BUCLES",
        ]),
    ],

    "2DO DE PRIMARIA": [
        ("UNIDAD 1 – ACTIVIDADES DE TECLADO Y MOUSE", [
            "CLASE 1.1 – HABILIDADES CON EL TECLADO – RAPIDTYPING",
            "CLASE 1.2 – HABILIDADES CON EL TECLADO – RAPIDTYPING",
            "CLASE 1.3 – ACTIVIDADES DE PRECISIÓN CON LOS CLICS",
            "CLASE 1.4 – ACTIVIDADES DE DRAG AND DROP",
        ]),
        ("UNIDAD 2 – LA COMPUTADORA Y SUS PARTES", [
            "CLASE 5.1 – LA COMPUTADORA",
            "CLASE 5.2 – HARDWARE Y SOFTWARE",
            "CLASE 5.3 – QUÉ SON LAS TICS",
            "CLASE 5.4 – QUÉ ES LA INTELIGENCIA ARTIFICIAL",
            "CLASE 5.5 – EVALUACIÓN DEL MÓDULO",
        ]),
        ("UNIDAD 3 – ALGORITMOS", [
            "CLASE 1.1 – REPASAMOS ALGORITMOS",
            "CLASE 1.2 – PRACTICANDO ALGORITMOS",
            "CLASE 1.3 – REPASAMOS BUCLES",
            "CLASE 1.4 – PRACTICAMOS BUCLES",
            "CLASE 1.5 – EVALUACIÓN DEL MÓDULO",
        ]),
        ("UNIDAD 4 – PROGRAMAS", [
            "CLASE 2.1 – ESCRIBIENDO PROGRAMAS",
            "CLASE 2.2 – PROGRAMAMOS CON BUCLES",
            "CLASE 2.3 – EVENTOS",
            "CLASE 2.4 – EVALUACIÓN MÓDULO 2",
        ]),
        ("UNIDAD 5 – CIUDADANÍA DIGITAL", [
            "CLASE 6.1 – QUÉ ES LA TECNOLOGÍA",
            "CLASE 6.2 – SEGURIDAD EN LÍNEA",
            "CLASE 6.3 – CONTRASEÑAS",
            "CLASE 6.4 – CIBERBULLYING",
        ]),
    ],

    "3RO DE PRIMARIA": [
        ("UNIDAD 1 – ACTIVIDADES DE TECLADO Y MOUSE", [
            "CLASE 1.1 – HABILIDADES CON EL TECLADO – RAPIDTYPING",
            "CLASE 1.2 – HABILIDADES CON EL TECLADO – RAPIDTYPING",
            "CLASE 1.3 – ACTIVIDADES DE PRECISIÓN CON LOS CLICS",
            "CLASE 1.4 – ACTIVIDADES DE DRAG AND DROP",
        ]),
        ("UNIDAD 2 – SISTEMAS COMPUTACIONALES", [
            "CLASE 1.1 – INTERNET",
            "CLASE 1.2 – NAVEGANDO EN LA WEB",
            "CLASE 1.3 – SEGURIDAD EN LÍNEA",
            "CLASE 1.4 – USO RESPONSABLE DE INTERNET",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – PROCESADORES DE TEXTO", [
            "CLASE 1.1 – PROCESADORES DE TEXTO – GOOGLE DOCS",
            "CLASE 1.2 – DISEÑO DE PÁGINA, FUENTES",
            "CLASE 1.3 – ESTILO, TAMAÑO Y COLOR DE FUENTES",
            "CLASE 1.4 – ALINEACIÓN DE TEXTO Y LISTAS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 4 – INTRODUCCIÓN A SPRITE LAB", [
            "CLASE 3.1 – INTRODUCCIÓN A SPRITE LAB",
            "CLASE 3.2 – COORDENADAS",
            "CLASE 3.3 – EVENTOS Y MOVIMIENTOS",
            "CLASE 3.4 – TEXTO EN SPRITE",
            "CLASE 3.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – VIDEOJUEGOS EN SPRITE LAB (NIVEL INICIAL)", [
            "CLASE 4.1 – METAMORFOSIS",
            "CLASE 4.2 – MASCOTA VIRTUAL",
            "CLASE 4.3 – ATRÁPAME SI PUEDES – CUESTIONARIO",
            "CLASE 4.4 – CLICKER GAMES",
            "CLASE 4.5 – PROYECTO DE CIERRE",
            "CLASE 4.6 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "4TO DE PRIMARIA": [
        ("UNIDAD 1 – ALGORITMOS", [
            "CLASE 1.1 – ALGORITMOS",
            "CLASE 1.2 – PROGRAMACIÓN EN CODEORG",
            "CLASE 1.3 – INTRODUCCIÓN A SCRATCH",
            "CLASE 1.4 – COORDENADAS",
            "CLASE 1.5 – PROYECTO DE CIERRE",
            "CLASE 1.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 2 – ANIMACIONES EN SCRATCH", [
            "CLASE 2.1 – BUCLES Y EVENTOS",
            "CLASE 2.2 – ANIMACIÓN DE UN COHETE",
            "CLASE 2.3 – ANIMACIÓN DE ANIMALES",
            "CLASE 2.4 – ANIMACIÓN DE SPRITES",
            "CLASE 2.5 – PROYECTO DE CIERRE",
            "CLASE 2.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 4 – HOJAS DE CÁLCULO NIVEL 1", [
            "CLASE 1.1 – COLUMNAS, FILAS Y CELDAS",
            "CLASE 1.2 – AJUSTANDO DIMENSIONES DE UNA CELDA Y ESTILOS",
            "CLASE 1.3 – TABLAS Y ALINEACIÓN DE TEXTO",
            "CLASE 1.4 – TIPOS DE DATOS Y CASILLAS DE VERIFICACIÓN CON DATOS LÓGICOS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – GOOGLE DOCS NIVEL 1", [
            "CLASE 1.1 – PROCESADORES DE TEXTO – GOOGLE DOCS",
            "CLASE 1.2 – DISEÑO DE PÁGINA, FUENTES",
            "CLASE 1.3 – ESTILO, TAMAÑO Y COLOR DE FUENTES",
            "CLASE 1.4 – ALINEACIÓN DE TEXTO Y LISTAS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "5TO DE PRIMARIA": [
        ("UNIDAD 1 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 2 – GOOGLE DOCS NIVEL 2", [
            "CLASE 1.1 – MULTIMEDIA EN GOOGLE DOCS",
            "CLASE 1.2 – INSERTAR ARCHIVOS MULTIMEDIA",
            "CLASE 1.3 – GOOGLE DOCS ADD-ONS",
            "CLASE 1.4 – MÁS ADD-ONS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – CIUDADANÍA DIGITAL", [
            "CLASE 5.1 – CYBERBULLYING",
            "CLASE 5.2 – GROOMING",
            "CLASE 5.3 – REDES SOCIALES",
            "CLASE 5.4 – PROYECTO DE CIERRE",
            "CLASE 5.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 4 – INTRODUCCIÓN A LAS PRESENTACIONES", [
            "CLASE 4.1 – GOOGLE SLIDES",
            "CLASE 4.2 – DIAPOSITIVAS, ESTILOS Y DISEÑOS",
            "CLASE 4.3 – TÍTULOS Y SUBTÍTULOS",
            "CLASE 4.4 – AGREGANDO IMÁGENES",
            "CLASE 4.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – REPASO DE SCRATCH", [
            "CLASE 5.1 – REPASO SCRATCH ANIMACIONES",
            "CLASE 5.2 – REPASO SCRATCH MENSAJERÍA",
            "CLASE 5.3 – REPASO SCRATCH VARIABLES",
            "CLASE 5.4 – PROYECTO CIERRE",
            "CLASE 5.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 6 – PROCEDIMIENTOS EN SCRATCH", [
            "CLASE 2.1 – PROCEDIMIENTOS EN SCRATCH",
            "CLASE 2.2 – DIBUJANDO CUADRADOS",
            "CLASE 2.3 – FLAPPY BIRDS",
            "CLASE 2.4 – PROYECTO DE CIERRE",
            "CLASE 2.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 7 – HOJAS DE CÁLCULO NIVEL 2", [
            "CLASE 1.1 – AGREGAR FILAS Y COLUMNAS",
            "CLASE 1.2 – FÓRMULAS, OPERADORES E IMPORTAR DATOS DE OTRAS CELDAS",
            "CLASE 1.3 – FÓRMULAS: TOTAL, PROMEDIO, MÁXIMO Y MÍNIMO",
            "CLASE 1.4 – UNIR CELDAS, INGRESAR FECHAS, HORAS, MINUTOS Y CÁLCULOS VARIOS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "6TO DE PRIMARIA": [
        ("UNIDAD 1 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 2 – HOJAS DE CÁLCULO NIVEL 2", [
            "CLASE 1.1 – MULTIMEDIA EN GOOGLE DOCS",
            "CLASE 1.2 – INSERTAR ARCHIVOS MULTIMEDIA",
            "CLASE 1.3 – GOOGLE DOCS ADD-ONS",
            "CLASE 1.4 – MÁS ADD-ONS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – GOOGLE FORMS", [
            "CLASE 1.1 – GOOGLE FORMS",
            "CLASE 1.2 – ENCUESTAS Y RECOPILACIÓN DE DATOS",
            "CLASE 1.3 – GOOGLE FORMS Y GOOGLE SHEETS",
            "CLASE 1.4 – CONFIGURACIONES",
            "CLASE 1.5 – PROYECTO INTEGRADOR",
        ]),
        ("UNIDAD 4 – INTRODUCCIÓN A ARCADE", [
            "CLASE 1.1 – INTRO A ARCADE",
            "CLASE 1.2 – SPRITES Y PIXEL ART",
            "CLASE 1.3 – VIDEOJUEGO CLICKERS",
            "CLASE 1.4 – COORDENADAS",
            "CLASE 1.5 – PROYECTO DE CIERRE",
            "CLASE 1.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – ARCADE: PROGRAMACIÓN DE VIDEOJUEGOS", [
            "CLASE 2.1 – ANIMACIONES",
            "CLASE 2.2 – GAMEPLAY",
            "CLASE 2.3 – BULLET HELL",
            "CLASE 2.4 – GAMEJAM",
            "CLASE 2.5 – PROYECTO DE CIERRE",
            "CLASE 2.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 6 – VIDEOJUEGOS TOPDOWN Y DE PLATAFORMAS", [
            "CLASE 3.1 – VIDEOJUEGOS TOP DOWN",
            "CLASE 3.2 – TILE MAPS Y LABERINTOS",
            "CLASE 3.3 – PROYECTO BIOMAS",
            "CLASE 3.4 – VIDEOJUEGO DE PLATAFORMAS",
            "CLASE 3.5 – PROYECTO DE CIERRE",
            "CLASE 3.6 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "1RO DE SECUNDARIA": [
        ("UNIDAD 1 – LAS COMPUTADORAS", [
            "CLASE 1.1 – COMPUTADORAS ¿PARA QUE SIRVEN?",
            "CLASE 1.2 – COMPUTADORAS, HISTORIA Y EVOLUCIÓN",
            "CLASE 1.3 – SOFTWARE Y HARDWARE",
            "CLASE 1.4 – INTERNET: NAVEGADORES Y SEGURIDAD EN LÍNEA",
            "CLASE 1.5 – PROYECTO DE CIERRE",
            "CLASE 1.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 2 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 3 – HOJAS DE CÁLCULO NIVEL 3", [
            "CLASE 1.1 – FÓRMULAS",
            "CLASE 1.2 – INSERTAR IMÁGENES",
            "CLASE 1.3 – HOJAS DE CÁLCULO COMPARTIDAS",
            "CLASE 1.4 – LISTAS DESPLEGABLES",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 4 – ALGORITMOS Y PROGRAMAS", [
            "CLASE 2.1 – ALGORITMOS",
            "CLASE 2.2 – PROGRAMAS ¿QUÉ SON?",
            "CLASE 2.3 – BOLITAS DE COLORES",
            "CLASE 2.4 – SACANDO BOLITAS – NUEVOS COMANDOS",
            "CLASE 2.5 – PROYECTO DE CIERRE",
            "CLASE 2.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – PROCEDIMIENTOS SIMPLES", [
            "CLASE 3.1 – PROCEDIMIENTOS",
            "CLASE 3.2 – ESTRATEGIAS Y PROCEDIMIENTOS",
            "CLASE 3.3 – AVENTURAS EN GOBSTONES",
            "CLASE 3.4 – ARMANDO PROCEDIMIENTOS",
            "CLASE 3.5 – PROYECTO DE CIERRE",
            "CLASE 3.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 6 – REPETICIONES SIMPLES", [
            "CLASE 4.1 – REPETICIONES SIMPLES",
            "CLASE 4.2 – PRACTICANDO REPETICIONES SIMPLES",
            "CLASE 4.3 – REPETICIONES ANIDADAS",
            "CLASE 4.4 – PROCEDIMIENTOS CON PARÁMETROS",
            "CLASE 4.5 – PROYECTO DE CIERRE",
            "CLASE 4.6 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "2DO DE SECUNDARIA": [
        ("UNIDAD 1 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 2 – GOOGLE FORMS", [
            "CLASE 1.1 – GOOGLE FORMS",
            "CLASE 1.2 – ENCUESTAS Y RECOPILACIÓN DE DATOS",
            "CLASE 1.3 – GOOGLE FORMS Y GOOGLE SHEETS",
            "CLASE 1.4 – CONFIGURACIONES",
            "CLASE 1.5 – PROYECTO INTEGRADOR",
        ]),
        ("UNIDAD 3 – HOJAS DE CÁLCULO NIVEL 4", [
            "CLASE 1.1 – ESTILOS CONDICIONALES",
            "CLASE 1.2 – GRÁFICOS",
            "CLASE 1.3 – MÁS GRÁFICOS",
            "CLASE 1.4 – LISTAS Y GRÁFICOS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 4 – INTRODUCCIÓN A PSEINT", [
            "CLASE 1.1 – ALGORITMOS Y PSEUDOCÓDIGO",
            "CLASE 1.2 – INTRODUCCIÓN A PSEINT",
            "CLASE 1.3 – VARIABLES EN PSEINT",
            "CLASE 1.4 – OPERADORES EN PSEINT",
            "CLASE 1.5 – PROYECTO DE CIERRE",
            "CLASE 1.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – ESTRUCTURAS DE CONTROL Y TOMA DE DECISIONES", [
            "CLASE 2.1 – INTRODUCCIÓN A LAS ESTRUCTURAS DE CONTROL",
            "CLASE 2.2 – ESTRUCTURAS SELECTIVAS DOBLES",
            "CLASE 2.3 – ESTRUCTURAS DE DECISIÓN COMPLEJAS: CONDICIONALES",
            "CLASE 2.4 – OPTIMIZACIÓN DE DECISIONES: USO DE «SEGÚN»",
            "CLASE 2.5 – PROYECTO DE CIERRE",
            "CLASE 2.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 6 – ESTRUCTURAS REPETITIVAS", [
            "CLASE 3.1 – ESTRUCTURAS REPETITIVAS",
            "CLASE 3.2 – ESTRUCTURAS REPETITIVAS «HACER» MIENTRAS",
            "CLASE 3.3 – BUCLE «PARA»",
            "CLASE 3.4 – BUCLES ANIDADOS",
            "CLASE 3.5 – PROYECTO DE CIERRE",
            "CLASE 3.6 – EVALUACIÓN MÓDULO",
        ]),
    ],

    "3RO DE SECUNDARIA": [
        ("UNIDAD 1 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 2 – HOJAS DE CÁLCULO NIVEL 5", [
            "CLASE 1.1 – REPASO DE LISTAS",
            "CLASE 1.2 – REPASO Y EJERCITACIÓN DE LISTAS Y GRÁFICOS",
            "CLASE 1.3 – FUNCIÓN IF",
            "CLASE 1.4 – FUNCIÓN SWITCH",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – INTRODUCCIÓN A PYTHON", [
            "CLASE 1.1 – LENGUAJES Y IDE «THONNY»",
            "CLASE 1.2 – VARIABLES Y OPERACIONES ARITMÉTICAS",
            "CLASE 1.3 – TIPOS DE DATOS Y FUNCIÓN «INPUT»",
            "CLASE 1.4 – PROYECTOS DE CIERRE",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 4 – CONDICIONALES, LISTAS Y TIPO DE DATO STRING", [
            "CLASE 2.1 – IF, ELSE, ELIF",
            "CLASE 2.2 – LISTAS",
            "CLASE 2.3 – TIPO DE DATO STRING CON CONDICIONALES",
            "CLASE 2.4 – RESOLUCIÓN DE PROBLEMAS PRÁCTICOS",
            "CLASE 2.5 – PROYECTO DE CIERRE",
            "CLASE 2.6 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – PROGRAMACIÓN CON ARDUINO", [
            "CLASE 1.1 – ¿QUÉ ES Y CÓMO FUNCIONA ARDUINO?",
            "CLASE 1.2 – DIFERENCIAS ENTRE ARDUINO Y RASPBERRY",
            "CLASE 1.3 – COMPONENTES DEL ARDUINO",
            "CLASE 1.4 – IDE",
            "CLASE 1.5 – PREPARANDO ARDUINO",
            "CLASE 1.6 – NUESTRO PRIMER PROGRAMA",
        ]),
        ("UNIDAD 6 – PROGRAMACIÓN FÍSICA", [
            "CLASE 2.1 – LEDS",
            "CLASE 2.2 – SENSORES DE LUMINOSIDAD",
            "CLASE 2.3 – PROGRAMACIÓN",
            "CLASE 2.4 – AND Y OR",
            "CLASE 2.5 – REPRESENTACIÓN DE DATOS EN FORMATO DE GRÁFICO",
        ]),
    ],

    "4TO Y 5TO DE SECUNDARIA": [
        ("UNIDAD 1 – GOOGLE DRIVE / GOOGLE CLASSROOM", [
            "CLASE 1.1 – GOOGLE DRIVE",
            "CLASE 1.2 – SUBIR Y DESCARGAR ARCHIVOS",
            "CLASE 1.3 – GOOGLE PHOTOS",
            "CLASE 1.4 – GOOGLE CLASSROOM",
            "CLASE 1.5 – LOG IN",
            "CLASE 1.6 – AULAS Y TAREAS",
            "CLASE 1.7 – ENTREGA DE PROYECTOS",
            "CLASE 1.8 – EVALUACIÓN DE MÓDULO",
        ]),
        ("UNIDAD 2 – HOJAS DE CÁLCULO NIVEL 5", [
            "CLASE 2.1 – REPASO DE LISTAS",
            "CLASE 2.2 – REPASO Y EJERCITACIÓN DE LISTAS Y GRÁFICOS",
            "CLASE 2.3 – FUNCIÓN IF",
            "CLASE 2.4 – FUNCIÓN SWITCH",
            "CLASE 2.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 3 – GENIALLY", [
            "CLASE 1.1 – GENIALLY",
            "CLASE 1.2 – INTERACTIVIDAD DE PRESENTACIONES",
            "CLASE 1.3 – ANIMACIONES",
            "CLASE 1.4 – GAMIFICACIÓN",
        ]),
        ("UNIDAD 4 – INTRODUCCIÓN A HTML", [
            "CLASE 1.1 – HIPERTEXTO, IDE",
            "CLASE 1.2 – ESTRUCTURAS",
            "CLASE 1.3 – ICONOS E IMÁGENES",
            "CLASE 1.4 – HTML LISTAS Y TABLAS",
            "CLASE 1.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 5 – HTML: FORMULARIOS Y MÚLTIPÁGINAS", [
            "CLASE 2.1 – FORMULARIOS EN HTML",
            "CLASE 2.2 – ENLACES Y MÚLTIPÁGINAS",
            "CLASE 2.3 – PÁGINA WEB 1",
            "CLASE 2.4 – PÁGINA WEB 2",
            "CLASE 2.5 – EVALUACIÓN MÓDULO",
        ]),
        ("UNIDAD 6 – PROGRAMACIÓN CON ARDUINO", [
            "CLASE 1.1 – ¿QUÉ ES Y CÓMO FUNCIONA ARDUINO?",
            "CLASE 1.2 – DIFERENCIAS ENTRE ARDUINO Y RASPBERRY",
            "CLASE 1.3 – COMPONENTES DEL ARDUINO",
            "CLASE 1.4 – IDE",
            "CLASE 1.5 – PREPARANDO ARDUINO",
            "CLASE 1.6 – NUESTRO PRIMER PROGRAMA",
        ]),
        ("UNIDAD 7 – PROGRAMACIÓN FÍSICA", [
            "CLASE 2.1 – LEDS",
            "CLASE 2.2 – SENSORES DE LUMINOSIDAD",
            "CLASE 2.3 – PROGRAMACIÓN",
            "CLASE 2.4 – AND Y OR",
            "CLASE 2.5 – REPRESENTACIÓN DE DATOS EN FORMATO DE GRÁFICO",
        ]),
    ],
}

# Create document
doc = Document()

# Title
title = doc.add_heading("CONTENIDO DE CLASES – TECNOLOGÍA 2026", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("EDUCACIÓN PRIMARIA Y SECUNDARIA")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.bold = True
    run.font.size = Pt(14)

doc.add_paragraph()  # spacer

# Add each course
for nivel, unidades in cursos.items():
    # Course level heading
    h = doc.add_heading(f"TECNOLOGÍA – {nivel}", 1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add units
    for unidad_nombre, clases in unidades:
        # Unit heading
        uh = doc.add_heading(unidad_nombre, 2)
        
        # Add classes as list
        for clase in clases:
            p = doc.add_paragraph(clase, style='List Bullet')
    
    doc.add_paragraph()  # spacer between levels

# Save
output_path = r"C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\LISTADO DE CONTENIDOS\CONTENIDOS DE TECNOLOGIA POR CURSO\CONTENIDO_TECNOLOGIA_2026_TODOS_LOS_NIVELES.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
