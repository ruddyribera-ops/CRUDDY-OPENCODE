# -*- coding: utf-8 -*-
"""Proper verification of Tutor 5P weekly plans against PDC."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os, re

BASE_PDC = r'C:\Users\Windows\Desktop\PDC - T2 - 2026\Tutor 5p'
WEEKLY = r'C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\5TO PRIMARIA _ TUTOR\5P'

def parse_pdc_weeks(path):
    doc = Document(path)
    t = doc.tables[1]
    content = t.rows[1].cells[1].text
    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]
    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]
    def find_unit(wpos):
        for i in range(len(units)-1, -1, -1):
            if units[i][0] < wpos:
                return 'UNIDAD %s: %s' % (units[i][1], units[i][2])
        return ''
    result = {}
    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        next_pos = weeks_raw[idx+1][0] if idx+1 < len(weeks_raw) else len(content)
        raw = content[wpos:next_pos]
        lines = [l.strip() for l in raw.strip().split('\n')]
        content_lines = [l for l in lines[1:] if l and not re.match(r'^[\*]+', l) and not re.match(r'^UNIDAD', l)]
        result[wnum] = {'unit': find_unit(wpos), 'content': '\n'.join(content_lines)}
    return result

# PDC files for each subject
PDCs = {
    'Lenguaje':    ('PDC_Lenguaje_5_Primaria.docx',    'LENGUAJE'),
    'Matematicas': ('PDC_Matematica_5_Primaria.docx',  'MATEMATICAS'),
    'Sociales':    ('PDC_Sociales_5_Primaria.docx',    'SOCIALES'),
    'Tecnologia':  ('PDC_Tecnologia_5_Primaria.docx',  'TECNOLOGIA'),
    'Valores':     ('PDC_Valores_5_Primaria.docx',     'VALORES'),
}

# Load all PDCs
pdc_data = {}
for subj, (fname, _) in PDCs.items():
    path = os.path.join(BASE_PDC, fname)
    if os.path.exists(path):
        pdc_data[subj] = parse_pdc_weeks(path)

# For each week, verify each subject that has a PDC
errors = []

# Check specific known cells from our earlier inspection
# The generator uses get_week_content() which correctly maps subjects to slots
# Let's verify by checking the actual cells that SHOULD have content

# Build ground truth: for each subject, which (slot_idx, day) pairs have that subject?
TIME_SLOTS = [
    ("07:55", "08:40", {"LUNES": "LENGUAJE",    "MARTES": "MATEMATICAS", "MIERCOLES": "MATEMATICAS", "JUEVES": "MUSICA",      "VIERNES": "MATEMATICAS"}),
    ("08:40", "09:25", {"LUNES": "ED.FISICA",   "MARTES": "MATEMATICAS", "MIERCOLES": "MATEMATICAS", "JUEVES": "SCIENCE",    "VIERNES": "MATEMATICAS"}),
    ("09:25", "10:10", {"LUNES": "INGLES",       "MARTES": "INGLES",       "MIERCOLES": "INGLES",       "JUEVES": "INGLES",      "VIERNES": "INGLES"}),
    ("10:10", "10:30", None),  # SNACK
    ("10:30", "11:15", {"LUNES": "LENGUAJE",    "MARTES": "ED.FISICA",   "MIERCOLES": "SOCIALES",     "JUEVES": "PORTUGUES",   "VIERNES": "LENGUAJE"}),
    ("11:15", "12:00", {"LUNES": "LENGUAJE",    "MARTES": "LENGUAJE",    "MIERCOLES": "SOCIALES",     "JUEVES": "MATEMATICAS", "VIERNES": "LENGUAJE"}),
    ("12:00", "12:15", None),  # PAUSE
    ("12:15", "13:00", {"LUNES": "SCIENCE",     "MARTES": "TECNOLOGIA",  "MIERCOLES": "ARTES",        "JUEVES": "ARTES",       "VIERNES": "HABILIDADES"}),
]
# slot_to_content_row: [2, 4, 6, 8, 10, 12, 14, 16]
slot_content_rows = [2, 4, 6, 8, 10, 12, 14, 16]
day_cols = {"LUNES": 1, "MARTES": 2, "MIERCOLES": 3, "JUEVES": 4, "VIERNES": 5}

# Build slot->(day->col) mapping for each subject we have PDC for
subject_slots = {}
for subj, (_, key) in PDCs.items():
    if not key:
        continue
    entries = []
    for si, slot in enumerate(TIME_SLOTS):
        if slot[2] is None:
            continue
        for day, subj_in_slot in slot[2].items():
            import unicodedata
            def norm(s):
                s = unicodedata.normalize('NFD', s)
                s = ''.join(c for c in s if c.isalnum()).upper()
                return s
            if norm(subj_in_slot) == norm(key):
                row = slot_content_rows[si]
                col = day_cols[day]
                entries.append((si, day, row, col))
    subject_slots[subj] = entries

print('Subject -> Slot/Row/Col mapping:')
for subj, entries in subject_slots.items():
    print('  %s: %s' % (subj, entries))

print()
print('=' * 70)
print('VERIFICATION: Checking ALL subject cells for weeks 15-30')
print('=' * 70)

total_checks = 0
total_errors = 0

for week in range(15, 31):
    weekly_path = os.path.join(WEEKLY, 'SEMANA %d - 5P.docx' % week)
    if not os.path.exists(weekly_path):
        errors.append('SEMANA %d: FILE MISSING' % week)
        continue

    doc = Document(weekly_path)
    t = doc.tables[0]

    for subj, entries in subject_slots.items():
        pdc_subj = pdc_data.get(subj, {})
        if not pdc_subj:
            continue

        for si, day, row, col in entries:
            total_checks += 1
            gen_text = t.rows[row].cells[col].text.strip() if row < len(t.rows) and col < len(t.rows[row].cells) else ''

            if week in [23, 24]:
                # Vacaciones - should say VACACIONES
                if gen_text and 'VACACIONES' not in gen_text:
                    errors.append('SEMANA %d %s %s: Expected VACACIONES, got: %s' % (week, subj, day, gen_text[:50]))
                    total_errors += 1
                continue

            if week not in pdc_subj:
                # Not in PDC for this week (e.g. Matematicas week 22)
                if gen_text:
                    errors.append('SEMANA %d %s %s: Have content but PDC missing. Content: %s' % (week, subj, day, gen_text[:50]))
                    total_errors += 1
                continue

            pdc_week = pdc_subj[week]
            pdc_content = pdc_week['content']
            pdc_unit = pdc_week['unit']

            # Check content is present
            pdc_first_line = pdc_content.split('\n')[0]

            # The generated text has "UNIDAD: ...\nCONTENIDO: ..." format
            if pdc_first_line and pdc_first_line[:25] not in gen_text[:100]:
                errors.append('SEMANA %d %s %s: CONTENT MISMATCH' % (week, subj, day))
                errors.append('  PDC first line: %s' % pdc_first_line[:80])
                errors.append('  GEN text:       %s' % gen_text[:80])
                total_errors += 1

print('Total checks: %d' % total_checks)
print('Total errors: %d' % total_errors)
if errors:
    for e in errors[:40]:
        print('  ERROR: ' + e)
    if len(errors) > 40:
        print('  ... and %d more errors' % (len(errors) - 40))
else:
    print('ALL CHECKS PASSED')