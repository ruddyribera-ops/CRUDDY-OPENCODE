# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os, re

BASE = r'C:\Users\Windows\Desktop\PDC - T2 - 2026\Tutor 5p'
WEEKLY = r'C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\5TO PRIMARIA _ TUTOR\5P'

def parse_pdc_weeks(path):
    doc = Document(path)
    t = doc.tables[1]
    content = t.rows[1].cells[1].text

    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]

    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]

    def find_unit(wpos):
        for i in range(len(units)-1, -1, -1):
            if units[i][0] < wpos:
                return 'UNIDAD %s: %s' % (units[i][1], units[i][2])
        return ''

    result = {}
    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        next_pos = weeks_raw[idx+1][0] if idx+1 < len(weeks_raw) else len(content)
        raw = content[wpos:next_pos]
        lines = [l.strip() for l in raw.strip().split('\n')]
        content_lines = [l for l in lines[1:] if l and not re.match(r'^[\*]+', l) and not re.match(r'^UNIDAD', l)]
        result[wnum] = {
            'unit': find_unit(wpos),
            'content': '\n'.join(content_lines),
            'header': lines[0]
        }
    return result


print('=' * 60)
print('CHECKING ALL SUBJECTS IN TUTOR 5P WEEKLY vs PDC')
print('=' * 60)

subjects = {
    'Lenguaje': 'PDC_Lenguaje_5_Primaria.docx',
    'Matematicas': 'PDC_Matematica_5_Primaria.docx',
    'Sociales': 'PDC_Sociales_5_Primaria.docx',
    'Tecnologia': 'PDC_Tecnologia_5_Primaria.docx',
    'Valores': 'PDC_Valores_5_Primaria.docx',
}

# Slot mapping: (slot_idx, day, table_row, table_col)
# Tecnologia slot=7, MARTES -> row=16, col=2
slot_map = {
    'Tecnologia': (7, 'MARTES', 16, 2),
}

for subj_name, fname in subjects.items():
    pdc_data = parse_pdc_weeks(os.path.join(BASE, fname))
    print('\n' + '=' * 60)
    print('SUBJECT:', subj_name)
    print('=' * 60)

    # Check weeks 15, 18, 22, 25, 30
    for week in [15, 16, 18, 22, 25, 30]:
        if week not in pdc_data:
            print('  Week %d: NOT IN PDC' % week)
            continue

        pdc_week = pdc_data[week]
        pdc_unit = pdc_week['unit']
        pdc_content = pdc_week['content']

        # Get generated doc
        doc = Document(os.path.join(WEEKLY, 'SEMANA %d - 5P.docx' % week))
        t = doc.tables[0]

        if subj_name in slot_map:
            slot_idx, day, row, col = slot_map[subj_name]
            gen_text = t.rows[row].cells[col].text
        else:
            # For other subjects just check if we can find them
            gen_text = '[SKIPPING - no slot map]'

        print('\n  Week %d:' % week)
        print('    PDC unit:    %s' % pdc_unit)
        print('    PDC content: %s' % pdc_content[:80])
        print('    GEN text:    %s' % gen_text[:80])

        # Simple match check
        pdc_short = pdc_content[:50].strip()
        gen_short = gen_text[:50].strip()
        if pdc_short == gen_short:
            print('    MATCH: YES')
        else:
            print('    MATCH: NO')
            print('    DIFF PDC:    "%s"' % pdc_short)
            print('    DIFF GEN:    "%s"' % gen_short)