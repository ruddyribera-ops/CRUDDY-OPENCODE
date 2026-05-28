# -*- coding: utf-8 -*-
"""Create dead-simple guide DOCX - no paths, no terminal."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Temp\opencode\GUIA-PARA-PROFESORES.docx"

NAVY = RGBColor(0, 51, 102)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
ORANGE = RGBColor(200, 80, 0)
WHITE = RGBColor(255, 255, 255)
GREEN = RGBColor(0, 120, 0)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)


def shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t.upper())
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY


def body(t, sz=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(t)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.color.rgb = DARK


def step(num, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"PASO {num}: {title.upper()}")
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = ORANGE

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK


def copy_box(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Light gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    r = p.add_run(t)
    r.font.size = Pt(8.5)
    r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(40, 40, 40)


def tip(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TIP: " + t)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY


def div():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("─" * 75)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(200, 200, 200)


def ok_box(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F5E9')
    pPr.append(shd)
    r = p.add_run("  " + t)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = GREEN


# ─── PORTADA ───
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(30)
r = t.add_run("GUIA RAPIDA")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = NAVY

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.paragraph_format.space_after = Pt(4)
r = s.add_run("De PDC a Planificaciones Semanales")
r.font.size = Pt(15)
r.font.color.rgb = DARK

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.paragraph_format.space_after = Pt(30)
r = s2.add_run("Para profesores — Sin conocimientos tecnicos")
r.font.size = Pt(10)
r.font.italic = True
r.font.color.rgb = GRAY

big = doc.add_paragraph()
big.alignment = WD_ALIGN_PARAGRAPH.CENTER
big.paragraph_format.space_after = Pt(6)
r = big.add_run("5 MINUTOS")
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = ORANGE

s3 = doc.add_paragraph()
s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
s3.paragraph_format.space_after = Pt(30)
r = s3.add_run("Tiempo estimado para generar TODAS tus planificaciones")
r.font.size = Pt(10)
r.font.color.rgb = GRAY

div()

# ─── QUE ES ESTO ───
h1("Que vamos a hacer")
body("Tus archivos PDC se convierten automaticamente en planificaciones semanales.")
body("Solo subes archivos al chat, pegas un texto, y la AI genera todo.", bold=True)
body("No necesitas saber programacion, usar terminal, ni nada tecnico.")

ok_box("Solo necesitas: archivos PDC + UNA PLANIFICACION SEMANAL VIEJA (como plantilla) + navegador")
doc.add_paragraph()

# ─── 4 PASOS ───
h1("4 Pasos")

step(1, "Sube 2 archivos al chat",
     "1. Abre ChatGPT (chatgpt.com) o Gemini (gemini.google.com)\n"
     "2. Busca el icono de adjuntar (clip o +)\n"
     "3. Sube estos archivos:\n\n"
     "   A) TUS ARCHIVOS PDC  (uno por materia/grado)\n"
     "   B) UNA PLANIFICACION SEMANAL VIEJA  (como plantilla del formato)\n\n"
     "   IMPORTANTE: Sin la planificacion vieja, la AI no sabra\n"
     "   que dias y horas tienes cada materia. Si no tienes una,\n"
     "   pidele a la AI que te sugiera basandose en tus PDC.")

tip("La planificacion vieja puede ser de cualquier mes anterior. Solo necesita tener tus horarios correctos.")

step(2, "Copia y pega este texto en el chat",
     "Selecciona todo el texto del recuadro gris de abajo.\n"
     "Ctrl+A para seleccionar todo, Ctrl+C para copiar, Ctrl+V para pegar.\n\n"
     "Luego envia el mensaje.")

copy_box("""HOLA. NECESITO QUE GENERES MIS PLANIFICACIONES SEMANALES.

ARRIBA TE HE ADJUNTADO:
- MIS ARCHIVOS PDC (Planes de Desarrollo Curricular)
- UNA PLANIFICACION SEMANAL VIEJA (como plantilla del formato)

REVISA QUE HAYAS RECIBIDO AMBOS ARCHIVOS.

MI ESCUELA:
- Nombre: [NOMBRE DE TU ESCUELA]
- Trimestre: 2 (11 de mayo al 28 de agosto de 2026)
- Director/a: [NOMBRE DEL DIRECTOR]

POR FAVOR:

1. REVISA que hayas recibido:
   a) mis archivos PDC
   b) la planificacion semanal vieja (como plantilla)
2. USA la planificacion vieja para saber:
   - Que dias y horas tiene cada materia
   - Cual es el formato exacto del documento
3. GENERA 16 documentos de planificacion semanal para cada materia/grado
   (Semanas 15 a 30 del Trimestre 2)
4. CADA documento debe tener:
   - Numero de semana y fechas correctas
   - Contenido de cada clase (extraido del PDC)
   - El mismo formato que la planificacion vieja
   - Semanas 23 y 24 marcadas como VACACIONES DE INVIERNO
5. ENTREGAME los archivos para descargar

ANTES DE GENERAR, CONFIRMAME:
- Que recibiste los archivos PDC
- Que recibiste la planificacion vieja (plantilla)
- Que entendiste el formato de dias y horas

GRACIAS.""")

tip("Si la AI no confirma que recibio tus archivos, no sigas. Sube los archivos de nuevo.")

step(3, "Espera la confirmacion de la AI",
     "La AI te va a decir:\n"
     "  - Cuantos archivos PDC recibio\n"
     "  - Cuantos documentos va a generar\n\n"
     "Si esta todo bien, dime:  SI, GENERALOS")

step(4, "Descarga los archivos",
     "La AI te va a dar los archivos para descargar.\n"
     "Descarga el ZIP o los archivos individuales.\n"
     "Descomprime el ZIP y listo.")

div()

# ─── FECHAS ───
h1("Calendario del Trimestre 2")

tbl = doc.add_table(rows=17, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "Semana"
hdr[1].text = "Fechas"
for c in hdr:
    shd(c, "003366")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = WHITE

weeks = [
    ("15","11/05 - 15/05"),("16","18/05 - 22/05"),("17","25/05 - 29/05"),
    ("18","01/06 - 05/06"),("19","08/06 - 12/06"),("20","15/06 - 19/06"),
    ("21","22/06 - 26/06"),("22","29/06 - 03/07"),
    ("23","VACACIONES 07/07 - 10/07"),("24","VACACIONES 14/07 - 17/07"),
    ("25","20/07 - 24/07"),("26","27/07 - 31/07"),
    ("27","04/08 - 08/08"),("28","11/08 - 15/08"),("29","18/08 - 22/08"),("30","25/08 - 28/08"),
]
for ri, (w, d) in enumerate(weeks):
    row = tbl.rows[ri + 1]
    row.cells[0].text = w
    row.cells[1].text = d
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
        if w in ["23","24"]:
            shd(c, "FFD700")
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(8)

div()

# ─── FAQ ───
h1("Preguntas Frecuentes")

faqs = [
    ("No tengo una planificacion semanal vieja.",
     "Busca en tu correo o archivos cualquier planificacion de meses anteriores. Si no tienes ninguna, escribe: 'Generame la planificacion basandote en el horario tipico de Bolivia para [tu grado]'."),
    ("No tengo el icono de adjuntar.",
     "En ChatGPT: es un icono de clip o + al lado del cuadrito de texto. En Gemini: es un icono de + a la izquierda."),
    ("Mis archivos no se llaman como el ejemplo.",
     "Esta bien. Solo renombrarlos al formato:  Materia_Grado_Nivel.docx  Por ejemplo:  Ciencias_3_Secundaria.docx"),
    ("Tengo solo una materia.",
     "Funciona igual. Sube ese archivo PDC y la AI generara las 16 semanas."),
    ("La AI me da los archivos en partes.",
     "Normal. Pidele: 'Generame un archivo ZIP con todo' o 'Descargame todo en un archivo'."),
    ("No me dio opcion de descargar.",
     "Escribe: 'Dame el enlace de descarga' o 'Como descargo los archivos'."),
    ("Las planificaciones salieron con dias y horas incorrectos.",
     "Esto pasa porque no subiste la planificacion vieja como plantilla. Vuelve a empezar y adjunta tu planificacion vieja."),
]

for q, a in faqs:
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(6)
    pq.paragraph_format.space_after = Pt(1)
    rq = pq.add_run("P: " + q)
    rq.font.size = Pt(10)
    rq.font.bold = True
    rq.font.color.rgb = DARK

    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_after = Pt(6)
    ra = pa.add_run("R: " + a)
    ra.font.size = Pt(10)
    ra.font.color.rgb = GRAY

div()

# ─── FOOTER ───
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.paragraph_format.space_before = Pt(16)
r = f.add_run("Esta guia funciona con cualquier AI: ChatGPT, Gemini, Claude, Grok...")
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = GRAY

f2 = doc.add_paragraph()
f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = f2.add_run("Solo necesitas: archivos PDC + 5 minutos + navegador")
r2.font.size = Pt(9)
r2.font.bold = True
r2.font.color.rgb = NAVY

doc.save(OUT)
print(f"Guardado: {OUT}")