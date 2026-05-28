# -*- coding: utf-8 -*-
"""
Generador de planificaciones semanales de Tecnologia para T2 (Semanas 15-30)
11 grados x 16 semanas = 176 documentos
"""

import re
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_PDC = r"C:\Users\Windows\Desktop\PDC - T2 - 2026\Tecnologia"
BASE_WEEKLY = r"C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\TECNOLOGIA"

# ─── Horarios por grado (dia, hora_inicio, hora_fin) ───
SCHEDULES = {
    "1P": [("LUNES",      "08:40", "09:25")],
    "2P": [("LUNES",      "10:30", "11:15")],
    "3P": [("LUNES",      "12:10", "12:55")],
    "4P": [("MARTES",     "08:40", "09:25")],
    "6P": [("LUNES",      "12:15", "13:00")],
    "1S": [("LUNES",      "07:00", "07:45"), ("MIÉRCOLES", "07:00", "07:45")],
    "2S": [("LUNES",      "08:40", "09:25"), ("JUEVES",    "08:40", "09:25")],
    "3S": [("LUNES",      "09:40", "10:25"), ("MIÉRCOLES", "12:55", "13:40")],
    "4S": [("LUNES",      "12:55", "13:40"), ("MARTES",    "12:55", "13:40")],
    "5S": [("MIÉRCOLES",  "12:10", "12:55"), ("JUEVES",    "12:10", "12:55")],
    "6S": [("MARTES",     "07:50", "08:35"), ("JUEVES",    "07:50", "08:35")],
}

WEEK_DATES = {
    15: "11/05 - 15/05", 16: "18/05 - 22/05", 17: "25/05 - 29/05",
    18: "01/06 - 05/06", 19: "08/06 - 12/06", 20: "15/06 - 19/06",
    21: "22/06 - 26/06", 22: "29/06 - 03/07",
    23: "07/07 - 10/07", 24: "14/07 - 17/07",
    25: "20/07 - 24/07", 26: "27/07 - 31/07",
    27: "04/08 - 08/08", 28: "11/08 - 15/08", 29: "18/08 - 22/08", 30: "25/08 - 28/08",
}


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="003366", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for bname in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def parse_pdc(pdc_path):
    """Parse PDC and return dict of week_num -> {unit, content, header}."""
    doc = Document(pdc_path)
    table = doc.tables[1]
    content = table.rows[1].cells[1].text

    # Find all units and their positions
    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]

    # Find all weeks and their positions
    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]

    def find_unit(wpos):
        for i in range(len(units) - 1, -1, -1):
            if units[i][0] < wpos:
                return f"UNIDAD {units[i][1]}: {units[i][2]}"
        return ''

    result = {}

    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        # Find content between this week and next
        next_pos = weeks_raw[idx + 1][0] if idx + 1 < len(weeks_raw) else len(content)
        raw_text = content[wpos:next_pos]

        # Parse header and content from raw_text
        lines = raw_text.strip().split('\n')
        header = lines[0].strip() if lines else f'Semana {wnum} ({wdates})'

        # Content lines (skip week header line)
        content_lines = []
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            if re.match(r'^[\*]+', line):
                continue  # skip *** VACACIONES *** lines
            if re.match(r'^UNIDAD \d+:', line):
                continue  # skip unit headers
            content_lines.append(line)

        result[wnum] = {
            'header': header,
            'unit': find_unit(wpos),
            'content': '\n'.join(content_lines),
            'dates': wdates
        }

    # Add vacaciones weeks
    for w in [23, 24]:
        result[w] = {
            'header': f'Semana {w} ({WEEK_DATES[w]})',
            'unit': 'VACACIONES',
            'content': '*** VACACIONES DE INVIERNO ***\nSemana 23 y 24 (Julio 7-17, 2026)\n\nDescanso reactivador!',
            'dates': WEEK_DATES[w]
        }

    return result


def get_pdc_path(grade):
    if grade in ["1P", "2P", "3P", "4P", "6P"]:
        nivel = "PRIMARIA"
    else:
        nivel = "SECUNDARIA"

    grade_map = {
        "1P": "1_Primaria", "2P": "2_Primaria", "3P": "3_Primaria",
        "4P": "4_Primaria", "6P": "6_Primaria",
        "1S": "1_Secundaria", "2S": "2_Secundaria", "3S": "3_Secundaria",
        "4S": "4_Secundaria", "5S": "5_Secundaria", "6S": "6_Secundaria"
    }
    fname = f"PDC_Tecnologia_{grade_map[grade]}.docx"
    return os.path.join(BASE_PDC, nivel, fname)


def create_doc(grade, week, week_data, schedule):
    n_days = len(schedule)
    doc = Document()

    # Landscape page
    section = doc.sections[0]
    section.page_width  = Cm(35.56)
    section.page_height = Cm(21.59)
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"WEEK PLANNER {grade}-2026 | TRIMESTRE 2 | SEMANA {week} ({WEEK_DATES[week]})")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Table: 3 rows x (1 + n_days) cols
    n_cols = 1 + n_days
    tbl = doc.add_table(rows=3, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: HORA | DAY1 | DAY2 | ...
    row0 = tbl.rows[0]
    row0.cells[0].text = "HORA"
    set_cell_bg(row0.cells[0], "003366")
    p = row0.cells[0].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (day, hi, hf) in enumerate(schedule):
        cell = row0.cells[i + 1]
        cell.text = f"{day}\n{WEEK_DATES[week]}"
        set_cell_bg(cell, "003366")
        p = cell.paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 1: time | Tecnologia | ...
    row1 = tbl.rows[1]
    time_str = f"{schedule[0][1]} – {schedule[0][2]}"
    row1.cells[0].text = time_str
    set_cell_bg(row1.cells[0], "E6F2FF")
    p = row1.cells[0].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(n_days):
        cell = row1.cells[i + 1]
        cell.text = "Tecnologia"
        set_cell_bg(cell, "E6F2FF")
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 2: content
    row2 = tbl.rows[2]

    if week in [23, 24]:
        row2.cells[0].text = week_data['content']
        set_cell_bg(row2.cells[0], "FFD700")
        p = row2.cells[0].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j in range(1, n_cols):
            row2.cells[j].text = ""
    else:
        unit = week_data.get('unit', '')
        content = week_data.get('content', '')
        if unit == 'VACACIONES':
            full = content
        elif unit:
            full = f"{unit}\n{content}"
        else:
            full = content

        for j in range(n_cols):
            cell = row2.cells[j]
            cell.text = full
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)

    # Column widths
    for i, col in enumerate(tbl.columns):
        for cell in col.cells:
            cell.width = Cm(2.5) if i == 0 else Cm(5.5)

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Unidad Educativa Las Palmas | Tecnologia | Trimestre 2 - 2026 | Ruddy Ribera")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path)


def generate_all():
    grades = ["1P", "2P", "3P", "4P", "6P", "1S", "2S", "3S", "4S", "5S", "6S"]

    total = 0
    errors = []

    for grade in grades:
        pdc_path = get_pdc_path(grade)
        if not os.path.exists(pdc_path):
            errors.append(f"PDC no encontrado: {pdc_path}")
            continue

        print(f"Procesando {grade}...", end=" ")

        try:
            weeks_data = parse_pdc(pdc_path)
        except Exception as e:
            errors.append(f"Error parseando PDC {grade}: {e}")
            print(f"ERROR: {e}")
            continue

        # Output dir
        if grade in ["1P", "2P", "3P", "4P", "6P"]:
            out_dir = os.path.join(BASE_WEEKLY, "PRIMARIA", grade)
        else:
            out_dir = os.path.join(BASE_WEEKLY, "SECUNDARIA", grade)
        ensure_dir(out_dir)

        schedule = SCHEDULES[grade]
        count = 0

        for week in range(15, 31):
            if week in weeks_data:
                wd = weeks_data[week]
            else:
                wd = {'unit': '', 'content': 'Contenido no disponible'}

            try:
                doc = create_doc(grade, week, wd, schedule)
                out_file = os.path.join(out_dir, f"SEMANA {week} - {grade}.docx")
                doc.save(out_file)
                count += 1
            except Exception as e:
                errors.append(f"Error SEMANA {week} {grade}: {e}")

        print(f"{count} docs")
        total += count

    print(f"\n{'='*50}")
    print(f"Total generados: {total} documentos")
    if errors:
        print(f"\nErrores ({len(errors)}):")
        for e in errors:
            print(f"  - {e}")
    else:
        print("Sin errores.")


if __name__ == "__main__":
    generate_all()