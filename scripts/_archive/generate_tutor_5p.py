# -*- coding: utf-8 -*-
"""
Generador de planificaciones semanales TUTOR 5P (Semanas 15-30)
Weekly planner combinado: todas las materias + Tecnologia
"""

import re
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_PDC_TUTOR = r"C:\Users\Windows\Desktop\PDC - T2 - 2026\Tutor 5p"
OUT_DIR = r"C:\Users\Windows\Desktop\01_Escuela\PLANIFICACION 2026\PLANIFICACIONES SEMANALES\5TO PRIMARIA _ TUTOR\5P"

WEEK_DATES = {
    15: "11/05 - 15/05", 16: "18/05 - 22/05", 17: "25/05 - 29/05",
    18: "01/06 - 05/06", 19: "08/06 - 12/06", 20: "15/06 - 19/06",
    21: "22/06 - 26/06", 22: "29/06 - 03/07",
    23: "07/07 - 10/07", 24: "14/07 - 17/07",
    25: "20/07 - 24/07", 26: "27/07 - 31/07",
    27: "04/08 - 08/08", 28: "11/08 - 15/08", 29: "18/08 - 22/08", 30: "25/08 - 28/08",
}

DAYS = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES"]

# ─── HORARIO 5P (orden de filas) ───
# Cada fila: (hora_inicio, hora_fin, {dia: materia})
TIME_SLOTS = [
    ("07:55", "08:40", {"LUNES": "LENGUAJE",    "MARTES": "MATEMÁTICAS", "MIÉRCOLES": "MATEMÁTICAS", "JUEVES": "MÚSICA",    "VIERNES": "MATEMÁTICAS"}),
    ("08:40", "09:25", {"LUNES": "ED. FÍSICA",  "MARTES": "MATEMÁTICAS", "MIÉRCOLES": "MATEMÁTICAS", "JUEVES": "SCIENCE",   "VIERNES": "MATEMÁTICAS"}),
    ("09:25", "10:10", {"LUNES": "INGLÉS",      "MARTES": "INGLÉS",      "MIÉRCOLES": "INGLÉS",      "JUEVES": "INGLÉS",     "VIERNES": "INGLÉS"}),
    ("10:10", "10:30", None),  # SNACK TIME
    ("10:30", "11:15", {"LUNES": "LENGUAJE",    "MARTES": "ED. FÍSICA",  "MIÉRCOLES": "SOCIALES",    "JUEVES": "PORTUGUÉS",  "VIERNES": "LENGUAJE"}),
    ("11:15", "12:00", {"LUNES": "LENGUAJE",    "MARTES": "LENGUAJE",    "MIÉRCOLES": "SOCIALES",    "JUEVES": "MATEMÁTICAS","VIERNES": "LENGUAJE"}),
    ("12:00", "12:15", None),  # ACTIVE PAUSE
    ("12:15", "13:00", {"LUNES": "SCIENCE",     "MARTES": "TECNOLOGÍA",  "MIÉRCOLES": "ARTES",       "JUEVES": "ARTES",      "VIERNES": "HABILIDADES"}),
]

SUBJECTS = ["LENGUAJE", "MATEMÁTICAS", "SOCIALES", "TECNOLOGÍA", "VALORES", "SCIENCE", "ED. FÍSICA", "MÚSICA", "PORTUGUÉS", "ARTES", "HABILIDADES", "INGLÉS"]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="AAAAAA", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for bname in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def parse_pdc(pdc_path, subject_name):
    """Parse PDC and return dict of week_num -> {unit, content}."""
    if not os.path.exists(pdc_path):
        return {}

    doc = Document(pdc_path)
    table = doc.tables[1]
    content = table.rows[1].cells[1].text

    unit_pat = re.compile(r'^UNIDAD (\d+): (.+)$', re.MULTILINE)
    units = [(m.start(), m.group(1), m.group(2)) for m in unit_pat.finditer(content)]

    week_pat = re.compile(r'^Semana (\d+) \((\d+/\d+ - \d+/\d+)\)$', re.MULTILINE)
    weeks_raw = [(m.start(), int(m.group(1)), m.group(2)) for m in week_pat.finditer(content)]

    def find_unit(wpos):
        for i in range(len(units) - 1, -1, -1):
            if units[i][0] < wpos:
                return f"UNIDAD {units[i][1]}: {units[i][2]}"
        return ''

    result = {}

    for idx, (wpos, wnum, wdates) in enumerate(weeks_raw):
        next_pos = weeks_raw[idx + 1][0] if idx + 1 < len(weeks_raw) else len(content)
        raw_text = content[wpos:next_pos]
        lines = raw_text.strip().split('\n')

        content_lines = []
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            if re.match(r'^[\*]+', line):
                continue
            if re.match(r'^UNIDAD \d+:', line):
                continue
            content_lines.append(line)

        result[wnum] = {
            'unit': find_unit(wpos),
            'content': '\n'.join(content_lines)
        }

    return result


def get_week_content(all_subjects_data, week, day, time_slot_idx):
    """Get content for a specific day/time slot in a week."""
    slot = TIME_SLOTS[time_slot_idx]
    if slot[2] is None:
        return None  # Snack or pause

    subject = slot[2].get(day)
    if not subject:
        return None

    # Normalize (remove accents, dots, spaces, uppercase)
    import unicodedata
    def norm(s):
        s = unicodedata.normalize('NFD', s)
        s = ''.join(c for c in s if c.isalnum()).upper()
        return s

    subject_norm = norm(subject)

    if week in [23, 24]:
        return "*** VACACIONES DE INVIERNO ***"

    for sub_name, sub_data in all_subjects_data.items():
        if norm(sub_name) == subject_norm:
            if week in sub_data:
                d = sub_data[week]
                unit = d.get('unit', '')
                content = d.get('content', '')
                if unit == 'VACACIONES':
                    return f"*** VACACIONES DE INVIERNO ***"
                elif unit:
                    return f"UNIDAD: {unit}\nCONTENIDO: {content}"
                else:
                    return f"CONTENIDO: {content}"
    return ""


def create_tutor_doc(week, all_subjects_data):
    """Create a combined tutor weekly planner for 5P."""
    doc = Document()

    # Landscape
    section = doc.sections[0]
    section.page_width  = Cm(35.56)
    section.page_height = Cm(21.59)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"WEEK PLANNER 5P-2026 | TUTOR | TRIMESTRE 2 | SEMANA {week} ({WEEK_DATES[week]})")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Table: 17 rows x 6 cols (time + 5 days)
    tbl = doc.add_table(rows=17, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Header
    row0 = tbl.rows[0]
    row0.cells[0].text = "HORARIO"
    set_cell_bg(row0.cells[0], "003366")
    p = row0.cells[0].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, day in enumerate(DAYS):
        cell = row0.cells[i + 1]
        cell.text = f"{day}\n{WEEK_DATES[week]}"
        set_cell_bg(cell, "003366")
        p = cell.paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Time slot rows
    slot_row_start = 1
    for slot_idx, slot in enumerate(TIME_SLOTS):
        row_num = slot_row_start + slot_idx
        hi, hf, day_subj = slot

        # Time cell
        cell_time = tbl.rows[row_num].cells[0]
        cell_time.text = f"{hi} - {hf}"
        set_cell_bg(cell_time, "E6F2FF")
        p = cell_time.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if day_subj is None:
            # Special row (snack/pause)
            for i, day in enumerate(DAYS):
                cell = tbl.rows[row_num].cells[i + 1]
                cell.text = "SNACK TIME" if "10:30" in hi else "ACTIVE PAUSE"
                set_cell_bg(cell, "FFFACD")
                p = cell.paragraphs[0]
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Subject row
            for i, day in enumerate(DAYS):
                subject = day_subj.get(day, "")
                cell = tbl.rows[row_num].cells[i + 1]
                cell.text = subject
                # Color coding
                if subject == "TECNOLOGÍA":
                    set_cell_bg(cell, "CCE5FF")
                elif subject in ["LENGUAJE", "MATEMÁTICAS"]:
                    set_cell_bg(cell, "E6F2FF")
                elif subject in ["SOCIALES", "SCIENCE"]:
                    set_cell_bg(cell, "D4EDDA")
                elif subject in ["ED. FÍSICA", "MÚSICA", "ARTES", "HABILIDADES"]:
                    set_cell_bg(cell, "F8D7DA")
                elif subject == "PORTUGUÉS":
                    set_cell_bg(cell, "D1ECF1")
                elif subject == "VALORES":
                    set_cell_bg(cell, "FFF3CD")
                elif subject == "INGLÉS":
                    set_cell_bg(cell, "E2E3E5")
                else:
                    set_cell_bg(cell, "F8F9FA")
                p = cell.paragraphs[0]
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Content row (next row in table)
            content_row_num = slot_row_start + len(TIME_SLOTS) + slot_idx + 1
            if content_row_num < 17:
                # Adjust for snack and pause rows
                # We need to figure out actual table row
                pass

    # Rebuild with proper 2-row-per-slot structure (time+subject | content)
    # Actually, let's restructure as: 17 rows = [header] + 8 slots * 2 rows
    # But existing format has 17 rows with gaps... let me re-check

    doc2 = Document()
    section2 = doc2.sections[0]
    section2.page_width  = Cm(35.56)
    section2.page_height = Cm(21.59)
    section2.left_margin   = Cm(1.0)
    section2.right_margin  = Cm(1.0)
    section2.top_margin    = Cm(1.0)
    section2.bottom_margin = Cm(0.8)

    title2 = doc2.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = title2.add_run(f"WEEK PLANNER 5P-2026 | TUTOR | TRIMESTRE 2 | SEMANA {week} ({WEEK_DATES[week]})")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0, 51, 102)

    doc2.add_paragraph()

    # 17 rows: row0=header, rows 1-16 = time slots (some merged visual rows)
    tbl2 = doc2.add_table(rows=17, cols=6)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Header
    r0 = tbl2.rows[0]
    r0.cells[0].text = "HORARIO"
    set_cell_bg(r0.cells[0], "003366")
    p = r0.cells[0].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, day in enumerate(DAYS):
        cell = r0.cells[i + 1]
        cell.text = f"{day}\n{WEEK_DATES[week]}"
        set_cell_bg(cell, "003366")
        p = cell.paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Map TIME_SLOTS to actual table rows
    # We have 8 time slots, but we need to fit into rows 1-16 (15 rows)
    # Some slots share rows (snack/pause are in the middle)
    # Row layout: 1,2 | 3,4 | 5 | 6 | 7,8 | 9,10 | 11 | 12,13 | 14,15 | 16,17 ??

    # Let's just use the exact pattern from the existing doc:
    # Row 1: 07:55-08:40
    # Row 2: content
    # Row 3: 08:40-09:25
    # Row 4: content
    # Row 5: 09:25-10:10
    # Row 6: content (none -ingles all same)
    # Row 7: 10:10-10:30 (snack)
    # Row 8: content (none)
    # Row 9: 10:30-11:15
    # Row 10: content
    # Row 11: 11:15-12:00
    # Row 12: content
    # Row 13: 12:00-12:15 (pause)
    # Row 14: content (none)
    # Row 15: 12:15-13:00
    # Row 16: content

    # That's 16 rows (1-16) for content + 1 header = 17 total. Matches!

    slot_to_row = [1, 3, 5, 7, 9, 11, 13, 15]  # Which row has the time
    content_to_row = [2, 4, 6, 8, 10, 12, 14, 16]  # Which row has the content

    for slot_idx, slot in enumerate(TIME_SLOTS):
        hi, hf, day_subj = slot
        time_row = slot_to_row[slot_idx]
        content_row = content_to_row[slot_idx]

        # Time cell
        tbl2.rows[time_row].cells[0].text = f"{hi} - {hf}"
        set_cell_bg(tbl2.rows[time_row].cells[0], "E6F2FF")
        p = tbl2.rows[time_row].cells[0].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if day_subj is None:
            # Snack (10:10-10:30) or pause (12:00-12:15)
            label = "SNACK TIME" if hi == "10:10" else "ACTIVE PAUSE"
            for i, day in enumerate(DAYS):
                cell = tbl2.rows[time_row].cells[i + 1]
                cell.text = label
                set_cell_bg(cell, "FFFACD")
                p = cell.paragraphs[0]
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Empty content row
            for i in range(6):
                tbl2.rows[content_row].cells[i].text = ""
        else:
            # Subject + content rows
            for i, day in enumerate(DAYS):
                subject = day_subj.get(day, "")
                # Time row: just subject name
                cell_t = tbl2.rows[time_row].cells[i + 1]
                cell_t.text = subject
                # Color coding
                if subject == "TECNOLOGÍA":
                    set_cell_bg(cell_t, "CCE5FF")
                elif subject in ["LENGUAJE", "MATEMÁTICAS"]:
                    set_cell_bg(cell_t, "E6F2FF")
                elif subject in ["SOCIALES", "SCIENCE"]:
                    set_cell_bg(cell_t, "D4EDDA")
                elif subject in ["ED. FÍSICA", "MÚSICA", "ARTES", "HABILIDADES"]:
                    set_cell_bg(cell_t, "F8D7DA")
                elif subject == "PORTUGUÉS":
                    set_cell_bg(cell_t, "D1ECF1")
                elif subject == "VALORES":
                    set_cell_bg(cell_t, "FFF3CD")
                elif subject == "INGLÉS":
                    set_cell_bg(cell_t, "E2E3E5")
                else:
                    set_cell_bg(cell_t, "F8F9FA")
                p = cell_t.paragraphs[0]
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Content row
                cell_c = tbl2.rows[content_row].cells[i + 1]
                week_content = get_week_content(all_subjects_data, week, day, slot_idx)

                if week_content is None:
                    cell_c.text = ""
                elif week_content == "*** VACACIONES DE INVIERNO ***":
                    cell_c.text = week_content
                    set_cell_bg(cell_c, "FFD700")
                    p = cell_c.paragraphs[0]
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(8)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell_c.text = week_content
                    set_cell_borders(cell_c, "AAAAAA", "4")
                    p = cell_c.paragraphs[0]
                    p.runs[0].font.size = Pt(7)

            # Time column for content row
            tbl2.rows[content_row].cells[0].text = ""

    # Set column widths
    for col_idx, col in enumerate(tbl2.columns):
        for cell in col.cells:
            if col_idx == 0:
                cell.width = Cm(2.2)
            else:
                cell.width = Cm(5.5)

    # Footer
    doc2.add_paragraph()
    foot = doc2.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Unidad Educativa Las Palmas | 5to Primaria Tutor | Trimestre 2 - 2026 | Ruddy Ribera")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(128, 128, 128)

    return doc2


def generate_tutor_5p():
    # Load all PDCs
    pdc_dir = BASE_PDC_TUTOR
    all_subjects = {}

    pdc_files = {
        "Lenguaje": "PDC_Lenguaje_5_Primaria.docx",
        "Matematicas": "PDC_Matematica_5_Primaria.docx",
        "Sociales": "PDC_Sociales_5_Primaria.docx",
        "Tecnologia": "PDC_Tecnologia_5_Primaria.docx",
        "Valores": "PDC_Valores_5_Primaria.docx",
    }

    for subj, fname in pdc_files.items():
        path = os.path.join(pdc_dir, fname)
        print(f"Cargando {subj}...")
        data = parse_pdc(path, subj)
        all_subjects[subj] = data
        print(f"  -> {len(data)} semanas cargadas")

    # Generate weeks 15-30
    total = 0
    for week in range(15, 31):
        try:
            doc = create_tutor_doc(week, all_subjects)
            out_file = os.path.join(OUT_DIR, f"SEMANA {week} - 5P.docx")
            doc.save(out_file)
            total += 1
            print(f"Generado: SEMANA {week} - 5P.docx")
        except Exception as e:
            print(f"Error en semana {week}: {e}")

    print(f"\nTotal: {total} documentos generados en:\n{OUT_DIR}")


if __name__ == "__main__":
    generate_tutor_5p()